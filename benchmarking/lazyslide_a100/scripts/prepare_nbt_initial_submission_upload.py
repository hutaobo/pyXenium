"""Prepare a clean Nature Biotechnology initial-submission upload directory."""

from __future__ import annotations

import csv
import html
import re
import shutil
import subprocess
from datetime import datetime, timezone
from pathlib import Path, PureWindowsPath

try:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.shared import Inches, Pt, RGBColor
    from docx.oxml.ns import qn
except Exception:  # pragma: no cover - optional DOCX fallback dependency
    WD_ALIGN_PARAGRAPH = None
    OxmlElement = None
    Inches = None
    Pt = None
    RGBColor = None
    qn = None


REPO_ROOT = Path(__file__).resolve().parents[3]
PACKAGE_ROOT = (
    REPO_ROOT
    / "docs"
    / "_static"
    / "tutorials"
    / "multimodal_histoseg_lazyslide_breast_wta"
    / "naturebiotech_package"
)
FINAL_DIR = PACKAGE_ROOT / "FINAL_SUBMISSION_NBT_20260513"
NATURE_DIR = FINAL_DIR / "Nature_Enhanced_Assets"
COMPOSITE_FIGURE_BASE = NATURE_DIR / "Figure_1_mTM_NBT_Brief_Composite"
UPLOAD_DIR = PACKAGE_ROOT / "NBT_INITIAL_SUBMISSION_UPLOAD_20260515_ONEFIGURE"
ADMIN_DIR = PACKAGE_ROOT / "NBT_INTERNAL_ADMIN_20260515"
FIGURE_DIR = UPLOAD_DIR / "Figures"
SOURCE_DATA_DIR = UPLOAD_DIR / "Source_Data"
SUPPLEMENTARY_DATA_DIR = UPLOAD_DIR / "Supplementary_Data"
SUPPLEMENTARY_DIR = UPLOAD_DIR / "Supplementary_Files"
ROBUSTNESS_SOURCE_DATA_DIR = REPO_ROOT / "manuscript" / "mtm_wta_nbt_replication" / "processed_data_archive_20260516" / "source_data"


def utc_now() -> str:
    return datetime.now(timezone.utc).replace(microsecond=0).isoformat()


def reset_upload_dir() -> None:
    resolved = UPLOAD_DIR.resolve()
    allowed_parent = PACKAGE_ROOT.resolve()
    if allowed_parent not in resolved.parents or not resolved.name.startswith("NBT_INITIAL_SUBMISSION_UPLOAD_"):
        raise RuntimeError(f"Refusing to reset unexpected directory: {resolved}")
    if resolved.exists():
        shutil.rmtree(resolved)
    FIGURE_DIR.mkdir(parents=True)
    SOURCE_DATA_DIR.mkdir(parents=True)
    SUPPLEMENTARY_DATA_DIR.mkdir(parents=True)
    SUPPLEMENTARY_DIR.mkdir(parents=True)


def reset_admin_dir() -> None:
    resolved = ADMIN_DIR.resolve()
    allowed_parent = PACKAGE_ROOT.resolve()
    if allowed_parent not in resolved.parents or not resolved.name.startswith("NBT_INTERNAL_ADMIN_"):
        raise RuntimeError(f"Refusing to reset unexpected directory: {resolved}")
    if resolved.exists():
        shutil.rmtree(resolved)
    ADMIN_DIR.mkdir(parents=True)


def copy_file(src: Path, dst: Path) -> None:
    if not src.exists():
        raise FileNotFoundError(src)
    dst.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(src, dst)


def copy_source_csv(
    src: Path,
    dst: Path,
    *,
    column_renames: dict[str, str] | None = None,
    value_replacements: dict[str, dict[str, str]] | None = None,
    basename_columns: set[str] | None = None,
) -> None:
    if not src.exists():
        raise FileNotFoundError(src)
    column_renames = column_renames or {}
    value_replacements = value_replacements or {}
    basename_columns = basename_columns or set()
    dst.parent.mkdir(parents=True, exist_ok=True)
    with src.open(newline="", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        if reader.fieldnames is None:
            raise ValueError(f"CSV has no header: {src}")
        fieldnames = [column_renames.get(name, name) for name in reader.fieldnames]
        with dst.open("w", newline="", encoding="utf-8") as out:
            writer = csv.DictWriter(out, fieldnames=fieldnames)
            writer.writeheader()
            for row in reader:
                cleaned: dict[str, str] = {}
                for old_name, value in row.items():
                    new_name = column_renames.get(old_name, old_name)
                    if old_name in basename_columns and value:
                        value = PureWindowsPath(value).name if "\\" in value else Path(value).name
                    if old_name in value_replacements:
                        value = value_replacements[old_name].get(value, value)
                    cleaned[new_name] = value
                writer.writerow(cleaned)


def extract_figure_legends(main_text: str) -> tuple[str, dict[str, str], str]:
    marker_match = re.search(r"^## Figure legends?\s*$", main_text, flags=re.MULTILINE)
    if marker_match is None:
        return main_text, {}, ""
    body = main_text[: marker_match.start()]
    legend_block = main_text[marker_match.end() :]
    declaration_match = re.search(r"^## Declarations\s*$", legend_block, flags=re.MULTILINE)
    declarations = ""
    if declaration_match is not None:
        declarations = legend_block[declaration_match.start() :].strip()
        legend_block = legend_block[: declaration_match.start()]
    matches = list(re.finditer(r"^### Figure ([1-4])\s*$", legend_block, flags=re.MULTILINE))
    legends: dict[str, str] = {}
    for idx, match in enumerate(matches):
        figure_id = match.group(1)
        start = match.end()
        end = matches[idx + 1].start() if idx + 1 < len(matches) else len(legend_block)
        legends[figure_id] = legend_block[start:end].strip()
    return body.rstrip(), legends, declarations


def copy_submission_figures() -> None:
    composite_files = [
        (COMPOSITE_FIGURE_BASE.with_suffix(".pdf"), FIGURE_DIR / "Figure_1.pdf"),
        (COMPOSITE_FIGURE_BASE.with_suffix(".png"), FIGURE_DIR / "Figure_1.png"),
    ]
    if all(src.exists() for src, _dst in composite_files):
        copies = composite_files
    else:
        copies = [
            (FINAL_DIR / "Figure_1_mTM_Framework.pdf", FIGURE_DIR / "Figure_1.pdf"),
            (FINAL_DIR / "Figure_1_mTM_Framework.png", FIGURE_DIR / "Figure_1.png"),
        ]
    supplementary_copies = [
        (
            FINAL_DIR / "Supplementary_Figure_SpatialPermutation_Defense.pdf",
            SUPPLEMENTARY_DIR / "Supplementary_Figure_1_SpatialPermutation_Defense.pdf",
        ),
        (
            NATURE_DIR / "Figure_2_Luminal_ER_HeroPatch_4Pairs_Nature.pdf",
            SUPPLEMENTARY_DIR / "Supplementary_Figure_2_HeroPatch_Examples.pdf",
        ),
        (
            NATURE_DIR / "Figure_2_Luminal_ER_HeroPatch_4Pairs_Nature.png",
            SUPPLEMENTARY_DIR / "Supplementary_Figure_2_HeroPatch_Examples.png",
        ),
        (
            FINAL_DIR / "Figure_3_MAZ_Stable_Coupled_Profiles.pdf",
            SUPPLEMENTARY_DIR / "Supplementary_Figure_3_MAZ.pdf",
        ),
        (
            FINAL_DIR / "Figure_3_MAZ_Stable_Coupled_Profiles.png",
            SUPPLEMENTARY_DIR / "Supplementary_Figure_3_MAZ.png",
        ),
    ]
    for src, dst in [*copies, *supplementary_copies]:
        copy_file(src, dst)


def copy_source_data() -> None:
    copy_source_csv(
        FINAL_DIR / "Source_Tables" / "Figure2_Luminal_ER_HeroPatch_4Pairs.csv",
        SOURCE_DATA_DIR / "Figure_1b_Hero_Patches_Source_Data.csv",
        column_renames={"patch_file": "patch_image"},
        basename_columns={"patch_file", "patch_image"},
    )
    copy_source_csv(
        FINAL_DIR / "Source_Tables" / "CrossCancer_Morphological_Signature_Table.csv",
        SOURCE_DATA_DIR / "Figure_1e_CrossCancer_Signature_Source_Data.csv",
    )
    copy_source_csv(
        FINAL_DIR / "Source_Tables" / "MAZ_QC_Table_v2.csv",
        SOURCE_DATA_DIR / "Figure_1d_MAZ_QC_Source_Data.csv",
        column_renames={"lead_lag_class": "boundary_profile_class"},
        value_replacements={
            "lead_lag_class": {
                "molecular_lead": "offset_boundary_zone",
                "morphology_lead": "offset_boundary_zone",
            }
        },
    )
    copy_source_csv(
        FINAL_DIR / "Source_Tables" / "Spatial_Permutation_Defense_Report.csv",
        SOURCE_DATA_DIR / "Figure_1c_Spatial_Permutation_Source_Data.csv",
    )
    copy_source_csv(
        FINAL_DIR / "Source_Tables" / "Spatial_BlockBootstrap_CI.csv",
        SOURCE_DATA_DIR / "Figure_1c_BlockBootstrap_Source_Data.csv",
    )
    supplementary_table_copies = {
        "Supplementary_Table_5_SpatialSensitivity_Source_Data.csv": "Supplementary_Table_5_SpatialSensitivity.csv",
        "Supplementary_Table_6_GeneComponent_Summary_Source_Data.csv": "Supplementary_Table_6_GeneComponent_Summary.csv",
        "Supplementary_Table_6_GeneComponent_Long_Source_Data.csv": "Supplementary_Table_6_GeneComponent_Long.csv",
        "Supplementary_Table_7_RegistrationPerturbation_Summary_Source_Data.csv": "Supplementary_Table_7_RegistrationPerturbation_Summary.csv",
        "Supplementary_Table_7_RegistrationPerturbation_Long_Source_Data.csv": "Supplementary_Table_7_RegistrationPerturbation_Long.csv",
        "Supplementary_Table_8_NestedSpatialHoldout_Summary_Source_Data.csv": "Supplementary_Table_8_NestedSpatialHoldout_Summary.csv",
        "Supplementary_Table_8_NestedSpatialHoldout_Long_Source_Data.csv": "Supplementary_Table_8_NestedSpatialHoldout_Long.csv",
    }
    for src_name, dst_name in supplementary_table_copies.items():
        copy_file(ROBUSTNESS_SOURCE_DATA_DIR / src_name, SUPPLEMENTARY_DATA_DIR / dst_name)


def make_manuscript_markdown() -> Path:
    main_text = (FINAL_DIR / "Main_Text.md").read_text(encoding="utf-8")
    body, legends, declarations = extract_figure_legends(main_text)
    composite_legend = legends.get("1", "")
    lines = [
        body,
        "",
        "## Figure legend",
        "",
        composite_legend,
        "",
    ]
    if declarations:
        lines.extend([declarations, ""])
    path = UPLOAD_DIR / "Manuscript_Initial_Submission.md"
    path.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")
    return path


def write_submission_guides() -> list[Path]:
    sig_path = UPLOAD_DIR / "SIGuide.md"
    sig_lines = [
        "# Supplementary Information Guide",
        "",
        "This guide maps the supplementary material submitted with the Nature Biotechnology Brief Communication initial package.",
        "",
        "## Primary supplementary document",
        "",
        "- `Supplementary_Information.docx`: supplementary methods, source-data mapping, spatial-null results, block-bootstrap intervals, program-family signatures, hero-patch metadata, spatial-sensitivity summaries, component-gene audits, registration-perturbation checks and nested spatial-holdout summaries.",
        "",
        "## Supplementary figure files",
        "",
        "- `Supplementary_Files/Supplementary_Figure_1_SpatialPermutation_Defense.pdf`: compartment-aware spatial-null defense.",
        "- `Supplementary_Files/Supplementary_Figure_2_HeroPatch_Examples.pdf`: breast S3 luminal estrogen-response hero-patch examples.",
        "- `Supplementary_Files/Supplementary_Figure_3_MAZ.pdf`: candidate molecularly active zone boundary profiles.",
        "",
        "## Source-data files",
        "",
        "- `Source_Data/Figure_1b_Hero_Patches_Source_Data.csv`",
        "- `Source_Data/Figure_1c_Spatial_Permutation_Source_Data.csv`",
        "- `Source_Data/Figure_1c_BlockBootstrap_Source_Data.csv`",
        "- `Source_Data/Figure_1d_MAZ_QC_Source_Data.csv`",
        "- `Source_Data/Figure_1e_CrossCancer_Signature_Source_Data.csv`",
        "",
        "## Supplementary data tables",
        "",
        "- `Supplementary_Data/Supplementary_Table_5_SpatialSensitivity.csv`",
        "- `Supplementary_Data/Supplementary_Table_6_GeneComponent_Summary.csv`",
        "- `Supplementary_Data/Supplementary_Table_6_GeneComponent_Long.csv`",
        "- `Supplementary_Data/Supplementary_Table_7_RegistrationPerturbation_Summary.csv`",
        "- `Supplementary_Data/Supplementary_Table_7_RegistrationPerturbation_Long.csv`",
        "- `Supplementary_Data/Supplementary_Table_8_NestedSpatialHoldout_Summary.csv`",
        "- `Supplementary_Data/Supplementary_Table_8_NestedSpatialHoldout_Long.csv`",
        "",
        "Raw 10x Genomics Atera WTA and H&E input files are not redistributed in this package. The public/vendor source datasets are available from 10x Genomics at https://www.10xgenomics.com/datasets/atera-wta-ffpe-human-breast-cancer and https://www.10xgenomics.com/datasets/atera-wta-ffpe-human-cervical-cancer.",
    ]
    sig_path.write_text("\n".join(sig_lines) + "\n", encoding="utf-8")

    reporting_path = UPLOAD_DIR / "Reporting_Summary_Draft.md"
    reporting_lines = [
        "# Reporting Summary Draft",
        "",
        "This draft is intended to populate the official Nature Portfolio Reporting Summary if the submission system requests the editable form. It is not a replacement for the official form.",
        "",
        "## Corresponding author",
        "",
        "Provided in the editorial cover letter for double-anonymous peer review.",
        "",
        "## Statistics",
        "",
        "- Exact sample sizes: primary breast S3 residual associations use 157 H&E-WTA paired contours; cervical stress-test associations use 215 H&E-WTA paired contours. Source-data tables report the contour count for each association.",
        "- Measurement unit: one statistical row is one registered spatial-omics contour with matched H&E embedding summaries and WTA program summaries.",
        "- Statistical tests: rank-residualized partial Spearman correlations were used for program-feature associations. Spatial-null empirical P values are two-sided and computed from compartment-aware permutations.",
        "- Covariates: spatial-omics-derived contour label where nonconstant, centroid-position covariates and boundary-distance summaries or bins. In breast S3 analyses the contour label is constant and omitted from the model matrix.",
        "- Multiple testing and discovery framing: the manuscript reports ranked candidate program-feature pairs and treats spatial permutations, block bootstrap, leave-one-block checks, local mismatch controls, centroid jitter, registration perturbation and nested spatial holdout as robustness checks rather than cohort-level validation.",
        "",
        "## Software and code",
        "",
        "- Data collection: no new experimental data collection was performed for this manuscript.",
        "- Data analysis: mTM analysis, direct WSI embedding, contour aggregation, statistical defense, figure composition and packaging are implemented in a reproducible Python workflow.",
        "- Code availability: code and package-release links are provided to editors in the cover letter and will be made public after double-anonymous review constraints are resolved.",
        "",
        "## Data",
        "",
        "- Source data: CSV source data for Fig. 1, supplementary figures and robustness summaries are included in the submission package.",
        "- Referenced datasets: raw 10x Genomics Atera WTA and H&E input files are public/vendor example datasets available at https://www.10xgenomics.com/datasets/atera-wta-ffpe-human-breast-cancer and https://www.10xgenomics.com/datasets/atera-wta-ffpe-human-cervical-cancer.",
        "- Restrictions: raw vendor inputs are not redistributed in the submission package.",
        "",
        "## Human research participants and data use",
        "",
        "- The analysis uses public/vendor example spatial WTA and H&E data.",
        "- The study did not involve new human sample collection, intervention, recruitment or generation of new identifiable participant data.",
        "- Sex, gender, race, ethnicity and other participant-level characteristics were not analyzed because the public/vendor example datasets do not support participant-level inference for this study.",
        "",
        "## Life sciences study design",
        "",
        "- Sample size rationale: all available paired contours passing the locked contour/WTA/image matching workflow were analyzed for the reported breast S3 and cervical stress-test association sets.",
        "- Data exclusions: contours without matched H&E-WTA summaries or required program/covariate values were not included in the corresponding association table; source-data files expose the analyzed contour counts.",
        "- Replication: cervical Atera WTA is treated as a second epithelial cancer stress test, not a direct replication of breast luminal estrogen-response biology. PLIP and UNI provide model-sensitivity evidence at the program-family level.",
        "- Randomization and blinding: no experimental randomization or blinding was introduced because this was a computational secondary analysis of public/vendor example data. Null distributions were generated by permutation within spatial/contour strata.",
        "- Validation limits: no IHC, protein-validation, clinical deployment or morphology-only diagnostic claim is made.",
    ]
    reporting_path.write_text("\n".join(reporting_lines) + "\n", encoding="utf-8")
    return [sig_path, reporting_path]


def inline_markdown_to_html(text: str) -> str:
    escaped = html.escape(text)
    escaped = re.sub(r"`([^`]+)`", r"<code>\1</code>", escaped)
    escaped = re.sub(r"\*\*([^*]+)\*\*", r"<strong>\1</strong>", escaped)
    escaped = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r'<a href="\2">\1</a>', escaped)
    escaped = escaped.replace("&lt;sup&gt;", "<sup>").replace("&lt;/sup&gt;", "</sup>")
    escaped = escaped.replace("&lt;sub&gt;", "<sub>").replace("&lt;/sub&gt;", "</sub>")
    return escaped


def flush_paragraph(buffer: list[str], parts: list[str]) -> None:
    if buffer:
        parts.append(f"<p>{inline_markdown_to_html(' '.join(buffer))}</p>")
        buffer.clear()


def markdown_to_html_document(markdown: str, title: str) -> str:
    parts = [
        "<!doctype html>",
        '<html><head><meta charset="utf-8">',
        f"<title>{html.escape(title)}</title>",
        "<style>",
        "body{font-family:Arial,Helvetica,sans-serif;font-size:11pt;line-height:1.35;color:#111;}",
        "h1{font-size:16pt;margin:0 0 10pt 0;} h2{font-size:13pt;margin:16pt 0 6pt 0;} h3{font-size:11.5pt;margin:12pt 0 4pt 0;}",
        "p{margin:0 0 8pt 0;} table{border-collapse:collapse;margin:8pt 0 12pt 0;width:100%;}",
        "th,td{border:1px solid #999;padding:4pt;vertical-align:top;} th{background:#f0f0f0;font-weight:bold;}",
        "code{font-family:Menlo,Consolas,monospace;font-size:9.5pt;} pre{font-family:Menlo,Consolas,monospace;font-size:9pt;white-space:pre-wrap;}",
        "</style></head><body>",
    ]
    lines = markdown.splitlines()
    paragraph: list[str] = []
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            flush_paragraph(paragraph, parts)
            i += 1
            continue
        if stripped == r"\[":
            flush_paragraph(paragraph, parts)
            math_lines: list[str] = []
            i += 1
            while i < len(lines) and lines[i].strip() != r"\]":
                math_lines.append(lines[i])
                i += 1
            parts.append(f"<pre>{html.escape(chr(10).join(math_lines))}</pre>")
            i += 1
            continue
        if stripped.startswith("|") and i + 1 < len(lines) and re.match(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$", lines[i + 1]):
            flush_paragraph(paragraph, parts)
            headers = [cell.strip() for cell in stripped.strip("|").split("|")]
            parts.append("<table><thead><tr>" + "".join(f"<th>{inline_markdown_to_html(cell)}</th>" for cell in headers) + "</tr></thead><tbody>")
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [cell.strip() for cell in lines[i].strip().strip("|").split("|")]
                parts.append("<tr>" + "".join(f"<td>{inline_markdown_to_html(cell)}</td>" for cell in cells) + "</tr>")
                i += 1
            parts.append("</tbody></table>")
            continue
        heading = re.match(r"^(#{1,3})\s+(.*)$", stripped)
        if heading:
            flush_paragraph(paragraph, parts)
            level = len(heading.group(1))
            parts.append(f"<h{level}>{inline_markdown_to_html(heading.group(2))}</h{level}>")
            i += 1
            continue
        if stripped.startswith("- "):
            flush_paragraph(paragraph, parts)
            parts.append("<ul>")
            while i < len(lines) and lines[i].strip().startswith("- "):
                parts.append(f"<li>{inline_markdown_to_html(lines[i].strip()[2:])}</li>")
                i += 1
            parts.append("</ul>")
            continue
        ordered = re.match(r"^\d+\.\s+(.*)$", stripped)
        if ordered:
            flush_paragraph(paragraph, parts)
            parts.append("<ol>")
            while i < len(lines):
                item_match = re.match(r"^\d+\.\s+(.*)$", lines[i].strip())
                if item_match is None:
                    break
                parts.append(f"<li>{inline_markdown_to_html(item_match.group(1))}</li>")
                i += 1
            parts.append("</ol>")
            continue
        paragraph.append(stripped)
        i += 1
    flush_paragraph(paragraph, parts)
    parts.append("</body></html>")
    return "\n".join(parts)


def clean_inline_text(text: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", lambda m: m.group(1) if m.group(1) == m.group(2) else f"{m.group(1)} ({m.group(2)})", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = text.replace("<sup>", "").replace("</sup>", "")
    text = text.replace("<sub>", "").replace("</sub>", "")
    return html.unescape(text)


def add_markdown_runs(
    paragraph,
    text: str,
    *,
    font_size: float | None = None,
    monospace_code: bool = True,
) -> None:
    token_re = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|<sup>.*?</sup>|<sub>.*?</sub>|\[[^\]]+\]\([^)]+\))")
    pos = 0
    for match in token_re.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(html.unescape(text[pos : match.start()]))
            if font_size is not None:
                run.font.size = Pt(font_size)
        token = match.group(0)
        run_text = clean_inline_text(token)
        run = paragraph.add_run(run_text)
        if token.startswith("**"):
            run.bold = True
        if token.startswith("`") and monospace_code:
            run.font.name = "Consolas"
        if token.startswith("<sup>"):
            run.font.superscript = True
        if token.startswith("<sub>"):
            run.font.subscript = True
        if font_size is not None:
            run.font.size = Pt(font_size)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(html.unescape(text[pos:]))
        if font_size is not None:
            run.font.size = Pt(font_size)


def remove_style_border(style) -> None:
    if qn is None:
        return
    p_pr = style._element.get_or_add_pPr()
    for border in list(p_pr.findall(qn("w:pBdr"))):
        p_pr.remove(border)


def set_style_font(style, *, name: str, size: float, bold: bool | None = None, color_rgb: tuple[int, int, int] = (0, 0, 0)) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    style.font.bold = bold
    if RGBColor is not None:
        style.font.color.rgb = RGBColor(*color_rgb)
    if qn is not None:
        r_pr = style._element.get_or_add_rPr()
        r_fonts = r_pr.rFonts
        if r_fonts is not None:
            for attr in ["ascii", "hAnsi", "eastAsia", "cs"]:
                r_fonts.set(qn(f"w:{attr}"), name)
            for attr in ["asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"]:
                r_fonts.attrib.pop(qn(f"w:{attr}"), None)
        color = r_pr.find(qn("w:color"))
        if color is not None:
            for attr in ["themeColor", "themeTint", "themeShade"]:
                color.attrib.pop(qn(f"w:{attr}"), None)


def set_document_defaults(document, *, profile: str = "compact") -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    if profile == "manuscript":
        font_name = "Times New Roman"
        normal_size = 12
        normal_spacing = 2.0
        normal_after = 6
        heading_specs = [
            ("Title", 16, 0, 12),
            ("Heading 1", 13, 12, 6),
            ("Heading 2", 12, 10, 4),
            ("Heading 3", 12, 8, 4),
        ]
    else:
        font_name = "Arial"
        normal_size = 10.5
        normal_spacing = 1.08
        normal_after = 7
        heading_specs = [
            ("Title", 18, 0, 5),
            ("Heading 1", 16, 10, 5),
            ("Heading 2", 13, 10, 5),
            ("Heading 3", 11.5, 10, 5),
        ]

    normal = document.styles["Normal"]
    set_style_font(normal, name=font_name, size=normal_size)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(normal_after)
    normal.paragraph_format.line_spacing = normal_spacing

    for style_name, size, before, after in heading_specs:
        try:
            style = document.styles[style_name]
        except KeyError:
            continue
        set_style_font(style, name=font_name, size=size, bold=True)
        remove_style_border(style)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15 if profile == "manuscript" else None

    for list_style_name in ["List Bullet", "List Number"]:
        try:
            style = document.styles[list_style_name]
        except KeyError:
            continue
        set_style_font(style, name=font_name, size=normal_size)
        style.paragraph_format.space_after = Pt(3 if profile == "manuscript" else 4)
        style.paragraph_format.line_spacing = 1.15 if profile == "manuscript" else 1.08

    for optional_style_name in ["Body Text", "First Paragraph"]:
        if optional_style_name in document.styles:
            style = document.styles[optional_style_name]
            set_style_font(style, name=font_name, size=normal_size)
            style.paragraph_format.space_before = Pt(0)
            style.paragraph_format.space_after = Pt(normal_after)
            style.paragraph_format.line_spacing = normal_spacing


def add_markdown_table(document, headers: list[str], rows: list[list[str]], *, monospace_code: bool = True) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    font_size = 7.2 if len(headers) >= 8 else 8.2 if len(headers) >= 5 else 9.2
    for cell, header in zip(table.rows[0].cells, headers, strict=False):
        paragraph = cell.paragraphs[0]
        add_markdown_runs(paragraph, header, font_size=font_size, monospace_code=monospace_code)
        for run in paragraph.runs:
            run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row, strict=False):
            paragraph = cell.paragraphs[0]
            add_markdown_runs(paragraph, value, font_size=font_size, monospace_code=monospace_code)
    document.add_paragraph()


def enable_continuous_line_numbers(document) -> None:
    if OxmlElement is None or qn is None:
        return
    for section in document.sections:
        sect_pr = section._sectPr
        for node in list(sect_pr.findall(qn("w:lnNumType"))):
            sect_pr.remove(node)
        line_numbers = OxmlElement("w:lnNumType")
        line_numbers.set(qn("w:countBy"), "1")
        line_numbers.set(qn("w:distance"), "360")
        line_numbers.set(qn("w:restart"), "continuous")
        sect_pr.append(line_numbers)


def add_centered_page_number_footer(document) -> None:
    if OxmlElement is None:
        return
    for section in document.sections:
        paragraph = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
        paragraph.clear()
        if WD_ALIGN_PARAGRAPH is not None:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        fld_separate = OxmlElement("w:fldChar")
        fld_separate.set(qn("w:fldCharType"), "separate")
        text = OxmlElement("w:t")
        text.text = "1"
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.extend([fld_begin, instr, fld_separate, text, fld_end])


def apply_review_numbering(dst: Path, *, line_numbers: bool, page_numbers: bool) -> None:
    try:
        from docx import Document
    except Exception:
        return
    document = Document(dst)
    if line_numbers:
        enable_continuous_line_numbers(document)
    if page_numbers:
        add_centered_page_number_footer(document)
    document.save(dst)


def convert_with_python_docx(src: Path, dst: Path, *, profile: str = "compact") -> bool:
    try:
        from docx import Document
    except Exception:
        return False
    if Inches is None or Pt is None:
        return False

    markdown = src.read_text(encoding="utf-8")
    document = Document()
    set_document_defaults(document, profile=profile)
    monospace_code = profile != "manuscript"
    lines = markdown.splitlines()
    paragraph: list[str] = []

    def flush_docx_paragraph() -> None:
        if paragraph:
            p = document.add_paragraph()
            add_markdown_runs(p, " ".join(paragraph), monospace_code=monospace_code)
            paragraph.clear()

    i = 0
    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped:
            flush_docx_paragraph()
            i += 1
            continue
        if stripped == r"\[":
            flush_docx_paragraph()
            math_lines: list[str] = []
            i += 1
            while i < len(lines) and lines[i].strip() != r"\]":
                math_lines.append(lines[i])
                i += 1
            p = document.add_paragraph("\n".join(math_lines))
            for run in p.runs:
                if profile == "manuscript":
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)
                else:
                    run.font.name = "Consolas"
                    run.font.size = Pt(9)
            i += 1
            continue
        if stripped.startswith("|") and i + 1 < len(lines) and re.match(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$", lines[i + 1]):
            flush_docx_paragraph()
            headers = [cell.strip() for cell in stripped.strip("|").split("|")]
            i += 2
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append([cell.strip() for cell in lines[i].strip().strip("|").split("|")])
                i += 1
            add_markdown_table(document, headers, rows, monospace_code=monospace_code)
            continue
        heading = re.match(r"^(#{1,3})\s+(.*)$", stripped)
        if heading:
            flush_docx_paragraph()
            level = len(heading.group(1))
            text = clean_inline_text(heading.group(2))
            document.add_heading(text, 0 if level == 1 else level)
            i += 1
            continue
        if stripped.startswith("- "):
            flush_docx_paragraph()
            while i < len(lines) and lines[i].strip().startswith("- "):
                p = document.add_paragraph(style="List Bullet")
                add_markdown_runs(p, lines[i].strip()[2:], monospace_code=monospace_code)
                i += 1
            continue
        ordered = re.match(r"^\d+\.\s+(.*)$", stripped)
        if ordered:
            flush_docx_paragraph()
            while i < len(lines):
                item_match = re.match(r"^\d+\.\s+(.*)$", lines[i].strip())
                if item_match is None:
                    break
                p = document.add_paragraph(style="List Number")
                add_markdown_runs(p, item_match.group(1), monospace_code=monospace_code)
                i += 1
            continue
        paragraph.append(stripped)
        i += 1
    flush_docx_paragraph()
    document.save(dst)
    return dst.exists()


def convert_with_textutil(src: Path, dst: Path) -> bool:
    html_path = dst.with_suffix(".html")
    html_path.write_text(markdown_to_html_document(src.read_text(encoding="utf-8"), src.stem), encoding="utf-8")
    try:
        subprocess.run(["textutil", "-convert", "docx", str(html_path), "-output", str(dst)], cwd=UPLOAD_DIR, check=True)
    except (FileNotFoundError, subprocess.CalledProcessError):
        return False
    finally:
        html_path.unlink(missing_ok=True)
    return dst.exists()


def make_pandoc_reference_docx(dst: Path, *, profile: str = "manuscript") -> Path | None:
    if dst.exists():
        return dst
    dst.parent.mkdir(parents=True, exist_ok=True)
    try:
        result = subprocess.run(
            ["pandoc", "--print-default-data-file", "reference.docx"],
            check=True,
            stdout=subprocess.PIPE,
        )
    except (OSError, subprocess.CalledProcessError):
        return None
    dst.write_bytes(result.stdout)
    try:
        from docx import Document
    except Exception:
        return dst
    if Inches is None or Pt is None:
        return dst
    document = Document(dst)
    set_document_defaults(document, profile=profile)
    document.save(dst)
    return dst


def convert_with_pandoc(src: Path, dst: Path, *, reference_doc: Path | None = None) -> bool:
    cmd = [
        "pandoc",
        "-f",
        "markdown+tex_math_dollars+tex_math_single_backslash",
        str(src),
        "-o",
        str(dst),
    ]
    if reference_doc is not None and reference_doc.exists():
        cmd.extend(["--reference-doc", str(reference_doc)])
    try:
        subprocess.run(cmd, cwd=UPLOAD_DIR, check=True)
    except (OSError, subprocess.CalledProcessError):
        return False
    return dst.exists()


def prepare_documents() -> None:
    copy_file(FINAL_DIR / "Cover_Letter.md", UPLOAD_DIR / "Cover_Letter.md")
    copy_file(FINAL_DIR / "Online_Methods.md", UPLOAD_DIR / "Online_Methods.md")
    copy_file(FINAL_DIR / "Supplementary_Information.md", UPLOAD_DIR / "Supplementary_Information.md")
    copy_file(FINAL_DIR / "Response_to_Reviewers.md", ADMIN_DIR / "Response_to_Likely_Reviewers_INTERNAL_NOT_FOR_INITIAL_UPLOAD.md")
    copy_file(FINAL_DIR / "Interactive_Hero_Contour_Browser.html", ADMIN_DIR / "Interactive_Hero_Contour_Browser_OPTIONAL.html")
    manuscript_md = make_manuscript_markdown()
    guide_docs = write_submission_guides()
    manuscript_reference_doc = make_pandoc_reference_docx(ADMIN_DIR / "Pandoc_Manuscript_Reference.docx", profile="manuscript")
    conversions = [
        (manuscript_md, UPLOAD_DIR / "Manuscript_Initial_Submission.docx"),
        (UPLOAD_DIR / "Online_Methods.md", UPLOAD_DIR / "Online_Methods.docx"),
        (UPLOAD_DIR / "Cover_Letter.md", UPLOAD_DIR / "Cover_Letter.docx"),
        (UPLOAD_DIR / "Supplementary_Information.md", UPLOAD_DIR / "Supplementary_Information.docx"),
        *[(path, path.with_suffix(".docx")) for path in guide_docs],
    ]
    conversion_lines = []
    for src, dst in conversions:
        manuscript_like_docs = {"Manuscript_Initial_Submission.md", "Online_Methods.md"}
        doc_profile = "manuscript" if src.name in manuscript_like_docs else "compact"
        if src.name == "Online_Methods.md" and convert_with_pandoc(src, dst, reference_doc=manuscript_reference_doc):
            status = "created_pandoc_docx_with_omml_math"
        elif convert_with_python_docx(src, dst, profile=doc_profile):
            status = "created_python_docx"
        elif convert_with_pandoc(src, dst, reference_doc=manuscript_reference_doc if doc_profile == "manuscript" else None):
            status = "created_pandoc"
        elif convert_with_textutil(src, dst):
            status = "created_textutil"
        else:
            status = "not_created"
        if status != "not_created":
            line_numbered_docs = {
                "Manuscript_Initial_Submission.md",
                "Online_Methods.md",
                "Supplementary_Information.md",
            }
            apply_review_numbering(dst, line_numbers=src.name in line_numbered_docs, page_numbers=True)
        conversion_lines.append(f"- `{dst.name}`: {status}")
    (ADMIN_DIR / "Pandoc_Conversion_Report.md").write_text("\n".join(["# Pandoc Conversion Report", "", *conversion_lines, ""]), encoding="utf-8")


def write_admin_readme() -> None:
    lines = [
        "# NBT Initial Submission Upload Package",
        "",
        f"Generated UTC: {utc_now()}",
        "",
        "## Upload first",
        "",
        "- `Manuscript_Initial_Submission.docx`: title page, abstract, main text, references, Fig. 1 caption and declarations.",
        "- `Online_Methods.docx`: online methods.",
        "- `Cover_Letter.docx`: cover letter.",
        "- `Supplementary_Information.docx`: supplementary information.",
        "- `SIGuide.docx`: supplementary-information guide for mapping uploaded supplementary files.",
        "- `Source_Data/`: CSV source data for quantitative figure panels.",
        "- `Supplementary_Data/`: CSV supplementary data tables for robustness and component-gene audits.",
        "",
        "## Complete the official form if requested",
        "",
        "- `Reporting_Summary_Draft.docx`: prefilled text for transferring into the official Nature Portfolio Reporting Summary form.",
        "",
        "## Upload if the submission system requests separate figure files",
        "",
        "- `Figures/Figure_1.pdf` and `Figures/Figure_1.png`: one composite main-text display item.",
        "- Legacy separate panels are in `Supplementary_Files/` for optional supplementary use, not as main-text display items.",
        "",
        "## Optional / internal",
        "",
        "- `Interactive_Hero_Contour_Browser_OPTIONAL.html` can be held back unless the journal accepts interactive supplementary files.",
        "- `Response_to_Likely_Reviewers_INTERNAL_NOT_FOR_INITIAL_UPLOAD.md` is an internal preparation file, not an initial-submission file.",
        "",
        "## What is deliberately not included",
        "",
        "- Hash manifests and provenance reports are not submission files. They remain in the working final package for internal reproducibility checks only.",
        "- Figure files do not contain standalone titles. Figure titles and interpretation are in the captions.",
        "",
        "## Guideline basis checked",
        "",
        "- Nature Biotechnology submission guidelines: https://www.nature.com/nbt/submission-guidelines",
        "- Nature Portfolio initial submission guidance: https://www.nature.com/nature/for-authors/initial-submission",
        "- Nature Research Figure Guide: https://research-figure-guide.nature.com/",
    ]
    (ADMIN_DIR / "UPLOAD_README_INTERNAL.md").write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    reset_upload_dir()
    reset_admin_dir()
    copy_submission_figures()
    copy_source_data()
    prepare_documents()
    write_admin_readme()
    print(f"Prepared upload package: {UPLOAD_DIR}")


if __name__ == "__main__":
    main()
