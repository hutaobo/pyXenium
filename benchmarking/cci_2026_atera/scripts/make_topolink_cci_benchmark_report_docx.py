"""Generate a formatted DOCX report for the TopoLink-CCI benchmark.

The report intentionally uses a curated document structure instead of directly
converting Markdown, because the benchmark summary contains wide tables and
multiple figure panels that need controlled placement.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[3]
REPORT_DIR = ROOT / "benchmarking" / "cci_2026_atera" / "reports"
RESULTS_DIR = ROOT / "benchmarking" / "cci_2026_atera" / "results"
OUTPUT_DOCX = REPORT_DIR / "topolink_cci_benchmark_manuscript_style_report.docx"


KEY_PATHS = [
    r"D:\GitHub\pyXenium\benchmarking\cci_2026_atera\results\publication_benchmark_24h_20260511",
    "/cfs/klemming/scratch/h/hutaobo/pyxenium_cci_benchmark_2026-04",
    "/data/taobo.hu/pyxenium_lr_benchmark_2026-04",
]


METHOD_ROWS = [
    ["Breast WTA", "TopoLink-CCI", "Full", "Main", "full common-db", "1,319,600", "PDC", "Primary clean full-common result."],
    ["Breast WTA", "CellPhoneDB", "Full", "Main", "full common-db", "reported in artifact", "PDC", "Reproducible non-spatial expression baseline."],
    ["Breast WTA", "LARIS", "Full", "Main", "full common-db", "reported in artifact", "PDC", "Diffusion/spatial smoothing comparator."],
    ["Breast WTA", "LIANA+", "Full", "Main", "full common-db", "reported in artifact", "PDC", "Spatial bivariate/multi-method comparator."],
    ["Breast WTA", "SpatialDM", "Full", "Main", "full common-db", "reported in artifact", "PDC", "Spatial co-expression/null comparator."],
    ["Breast WTA", "stLearn", "Full", "Main", "full common-db", "reported in artifact", "PDC/A100", "Spatial neighborhood permutation comparator."],
    ["Breast WTA", "Squidpy", "Full", "Supplement", "full common-db", "1,319,600", "A100", "Permutation-style ligrec supplement."],
    ["Breast WTA", "CellChat", "Full", "Supplement", "LR-only full", "31,592", "A100", "Pathway step disabled because of runtime bug."],
    ["Breast WTA", "COMMOT", "Full", "Appendix", "chunked full", "724,604", "A100/PDC", "Transport/diffusion comparator; 16/16 chunks merged."],
    ["Breast WTA", "Giotto", "Bounded", "Appendix", "pilot50k", "1,151,351", "A100", "Full limited by R Matrix/TsparseMatrix 2^31-1 sparse index boundary."],
    ["Breast WTA", "SpaTalk", "Bounded", "Appendix", "smoke20k", "37,532", "A100", "66/66 smoke chunks completed; pilot50k had no usable output."],
    ["Breast WTA", "NICHES", "Bounded", "Appendix", "pilot50k", "1,181,042", "A100", "Full stopped by memory safety gate."],
    ["Breast WTA", "CellNEST", "Bounded", "Appendix", "pilot50k", "157,553", "A100", "Bounded GPU/source workflow result."],
    ["Breast WTA", "CellAgentChat", "Bounded", "Appendix", "pilot50k", "1,143,543", "A100", "16/16 pilot chunks succeeded; full resource exceeded."],
    ["Breast WTA", "SCILD", "Bounded", "Bounded appendix", "smoke3k common-db", "2,880 rows", "A100", "Official source-backed ligand-diffusion bounded retry succeeded: 3000 cells x 20 CCI pairs, peak RSS about 18.9GB."],
    ["Breast WTA", "FastCCC", "Bounded", "Appendix", "smoke20k", "1,319,600", "A100", "A100 retry succeeded after PDC standardization failure; analytic non-spatial expression baseline."],
    ["Breast WTA", "Copulacci", "Failure card", "Terminal card", "env/API audit", "reported in artifact", "PDC/A100", "Official source imports, but no safe common-resource bounded adapter mapping."],
    ["Breast WTA", "NicheNet", "Failure card", "Terminal card", "env/API audit", "reported in artifact", "PDC/A100", "R dependency/API audit blocker; downstream support method, not direct spatial CCI ranker."],
    ["Cervical WTA", "TopoLink-CCI", "Full", "Cross-dataset", "full common-db", "2,404,971", "PDC", "Top hit: DSC2-DSG3 in differentiating tumor cells."],
]


FIGURES = [
    {
        "title": "Benchmark completion status",
        "path": RESULTS_DIR / "preview_completed_20260511" / "figures" / "fig1_completion_status.png",
        "source": RESULTS_DIR / "preview_completed_20260511" / "source_data" / "fig1_completion_status.tsv",
        "caption": "Overview of full, bounded, and reproducible failure-card method states.",
    },
    {
        "title": "Synthetic Truth validation",
        "path": RESULTS_DIR / "publication_benchmark_24h_20260511" / "pdc_collected" / "publication_benchmark_24h_20260511" / "synthetic_truth" / "figures" / "synthetic_truth_auroc_auprc.png",
        "source": RESULTS_DIR / "publication_benchmark_24h_20260511" / "pdc_collected" / "publication_benchmark_24h_20260511" / "synthetic_truth" / "tables" / "synthetic_truth_metrics.tsv",
        "caption": "TopoLink-CCI achieves AUROC 0.9919 and AUPRC 0.8333; topology-anchor-only loses precision-recall stability.",
    },
    {
        "title": "Breast WTA TopoLink-CCI top axes",
        "path": RESULTS_DIR / "preview_completed_20260511" / "figures" / "fig3_breast_topolink_top_axes.png",
        "source": RESULTS_DIR / "preview_completed_20260511" / "source_data" / "fig3_breast_topolink_top_axes.tsv",
        "caption": "High-ranking Breast WTA axes emphasize vascular, stromal, immune and Notch-associated CCI themes.",
    },
    {
        "title": "Canonical pair rank matrix",
        "path": RESULTS_DIR / "cross_method_comparison_20260511" / "figures" / "fig6_canonical_pair_rank_heatmap.png",
        "source": RESULTS_DIR / "cross_method_comparison_20260511" / "source_data" / "fig6_canonical_pair_rank_heatmap.tsv",
        "caption": "Canonical axis recovery is evaluated by within-method ranks rather than raw score comparison.",
    },
    {
        "title": "Runtime and scalability",
        "path": RESULTS_DIR / "cross_method_comparison_20260511" / "figures" / "fig1_runtime_hours.png",
        "source": RESULTS_DIR / "cross_method_comparison_20260511" / "source_data" / "fig1_runtime_hours.tsv",
        "caption": "Engineering comparison across full and bounded runs.",
    },
    {
        "title": "Top-100 pair Jaccard consistency",
        "path": RESULTS_DIR / "cross_method_comparison_20260511" / "figures" / "fig3_top100_pair_jaccard_heatmap.png",
        "source": RESULTS_DIR / "cross_method_comparison_20260511" / "source_data" / "fig3_top100_pair_jaccard_heatmap.tsv",
        "caption": "Cross-method consistency is summarized as overlap among top-ranked CCI axes.",
    },
    {
        "title": "Breast versus Cervical generalization",
        "path": RESULTS_DIR / "preview_completed_20260511" / "figures" / "fig4_breast_cervical_topolink_comparison.png",
        "source": RESULTS_DIR / "preview_completed_20260511" / "source_data" / "fig4_breast_cervical_topolink_comparison.tsv",
        "caption": "TopoLink-CCI recovers tissue-context-specific top axes across Breast and Cervical WTA datasets.",
    },
    {
        "title": "Cervical WTA TopoLink-CCI top axes",
        "path": RESULTS_DIR / "preview_completed_20260511" / "figures" / "fig3b_cervical_topolink_top_axes.png",
        "source": RESULTS_DIR / "preview_completed_20260511" / "source_data" / "fig3b_cervical_topolink_top_axes.tsv",
        "caption": "Cervical WTA full common-db result contains 2,404,971 rows and is led by DSC2-DSG3.",
    },
    {
        "title": "Result row counts",
        "path": RESULTS_DIR / "preview_completed_20260511" / "figures" / "fig2_result_row_counts.png",
        "source": RESULTS_DIR / "preview_completed_20260511" / "source_data" / "fig2_result_row_counts.tsv",
        "caption": "Full and bounded outputs differ in scale and should not be collapsed into a single evidence tier.",
    },
    {
        "title": "Overlap with TopoLink-CCI",
        "path": RESULTS_DIR / "cross_method_comparison_20260511" / "figures" / "fig5_overlap_with_topolink.png",
        "source": RESULTS_DIR / "cross_method_comparison_20260511" / "source_data" / "fig5_overlap_with_topolink.tsv",
        "caption": "Overlap with TopoLink-CCI is interpreted as theme-level support, not absolute ground truth.",
    },
    {
        "title": "Canonical recovery F1",
        "path": RESULTS_DIR / "cross_method_comparison_20260511" / "figures" / "fig2_canonical_recovery_f1.png",
        "source": RESULTS_DIR / "cross_method_comparison_20260511" / "source_data" / "fig2_canonical_recovery_f1.tsv",
        "caption": "F1 is used only for predefined canonical truth sets, not global real-tissue CCI ranking.",
    },
    {
        "title": "Exploratory ARI of top-k binary selections",
        "path": RESULTS_DIR / "cross_method_comparison_20260511" / "figures" / "fig4_top100_binary_selection_ari_heatmap.png",
        "source": RESULTS_DIR / "cross_method_comparison_20260511" / "source_data" / "fig4_top100_binary_selection_ari_heatmap.tsv",
        "caption": "ARI is shown as an exploratory top-k selection similarity metric, not a main endpoint.",
    },
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 7) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def set_document_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)

    styles = doc.styles
    for style_name, size, bold in [
        ("Normal", 9, False),
        ("Title", 20, True),
        ("Heading 1", 15, True),
        ("Heading 2", 12, True),
        ("Heading 3", 10, True),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = bold


def add_runs(paragraph, segments) -> None:
    for text, bold, italic in segments:
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = "Arial"
        run.font.size = Pt(9)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_paragraph(doc: Document, segments, style: str | None = None) -> None:
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.08
    add_runs(paragraph, segments)


def add_path_paragraph(doc: Document, prefix: str, path: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)
    add_runs(paragraph, [(prefix, False, False), (path, True, False)])


def add_method_table(doc: Document) -> None:
    headers = ["Dataset", "Method", "Status", "Evidence tier", "Phase", "Rows", "Hardware/source", "Interpretation"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_text(cell, header, bold=True, size=7)
        set_cell_shading(cell, "D9EAF7")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for row in METHOD_ROWS:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=6)
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.TOP
            if row[2] == "Full":
                set_cell_shading(cells[2], "DDEEDB")
            elif row[2] == "Bounded":
                set_cell_shading(cells[2], "FFF2CC")
            elif row[2] == "Failure card":
                set_cell_shading(cells[2], "F4CCCC")
            elif row[2] == "Deferred":
                set_cell_shading(cells[2], "E6E6E6")


def add_figure(doc: Document, figure: dict, index: int) -> None:
    doc.add_heading(f"Figure panel {index}. {figure['title']}", level=3)
    if figure["path"].exists():
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(figure["path"]), width=Inches(7.9))
    else:
        add_paragraph(doc, [(f"Missing figure file: {figure['path']}", True, False)])
    add_paragraph(doc, [(figure["caption"], False, False)])
    add_path_paragraph(doc, "Figure: ", str(figure["path"]))
    add_path_paragraph(doc, "Source data: ", str(figure["source"]))


def build_report() -> None:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    set_document_defaults(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_runs(title, [("TopoLink-CCI Benchmarking Manuscript-style Report", True, False)])

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_runs(subtitle, [("Formatted Word report with embedded figures and source-data links", False, True)])

    doc.add_heading("Overall Progress", level=1)
    add_paragraph(
        doc,
        [
            ("本报告整理 TopoLink-CCI benchmarking 的完整收口状态。主线从 PDC Synthetic Truth 和 false-positive controls 出发，说明方法在受控场景下的 ranking behavior，再引入 Breast WTA 主 benchmark、A100 supplement、bounded appendix 和 Cervical WTA cross-dataset generalization。", False, False)
        ],
    )
    add_paragraph(
        doc,
        [
            ("当前 PDC publication queue 为 ", False, False),
            ("0", True, False),
            ("，A100 无 active rescue process。报告口径已升级为 expanded 18-method benchmark：9 个 full success、7 个 bounded success、2 个 reproducible failure cards、0 个 deferred candidate methods、0 个 pending/running。这里的 all_methods_accounted=true 表示所有提到的方法都有明确状态，不表示所有方法都成功。", False, False),
        ],
    )
    add_paragraph(
        doc,
        [
            ("FastCCC、SCILD、Copulacci 和 NicheNet 现在正式计入 expanded denominator，并且均已在 PDC 与 A100 完成终端化处理。FastCCC 和 SCILD 经 A100 retry 升级为 bounded subset result；Copulacci 和 NicheNet 保留 reproducible failure card，原因分别是 adapter workflow mapping blocker 与 R dependency/API audit blocker。", False, False),
        ],
    )
    for p in KEY_PATHS:
        add_path_paragraph(doc, "Key path: ", p)

    doc.add_heading("Summary of Evaluated Methods", level=2)
    add_method_table(doc)
    add_figure(doc, FIGURES[0], 1)

    doc.add_page_break()
    doc.add_heading("Validation & Controls", level=1)
    add_paragraph(
        doc,
        [
            ("PDC Synthetic Truth 结果提供了真实 WTA benchmark 的前置可信度论证。完整 TopoLink-CCI 模型获得 ", False, False),
            ("AUROC 0.9919", True, False),
            (" 和 ", False, False),
            ("AUPRC 0.8333", True, False),
            ("。expression-only 为 AUROC 0.9839 / AUPRC 0.8333，contact-only 为 AUROC 0.9919 / AUPRC 0.8333，而 topology-anchor-only 虽然 AUROC 为 0.9839，但 AUPRC 降至 0.5833。", False, False),
        ],
    )
    add_paragraph(
        doc,
        [
            ("该结果说明 Topology Anchor 单独使用不足以稳定排序 positive CCI axes；完整模型通过 expression specificity 和 local contact 补充拓扑信息，提升 precision-recall ranking 的稳定性。", False, False)
        ],
    )
    add_figure(doc, FIGURES[1], 2)

    doc.add_page_break()
    doc.add_heading("Core Benchmarking Results", level=1)
    doc.add_heading("Breast WTA common-db main comparison", level=2)
    add_paragraph(
        doc,
        [
            ("Breast WTA 是主实验数据集。TopoLink-CCI Breast full result 产生 ", False, False),
            ("1,319,600 rows", True, False),
            ("，构成 biological interpretation、classic CCI axis selection、validation 和 cross-method consistency 的主要输入。", False, False),
        ],
    )
    add_path_paragraph(doc, "TopoLink-CCI Breast artifact: ", "pdc_collected/pdc_20260426_1327/runs/full_common/pyxenium")
    add_figure(doc, FIGURES[2], 3)
    add_figure(doc, FIGURES[3], 4)

    doc.add_heading("A100 supplementary full and appendix runs", level=2)
    add_paragraph(
        doc,
        [
            ("A100 supplement 提供 Squidpy full、CellChat LR-only full 和 COMMOT chunked result。bounded appendix 记录 Giotto、SpaTalk、NICHES、CellNEST 和 CellAgentChat 的 scalability/resource tradeoff。", False, False)
        ],
    )
    add_figure(doc, FIGURES[4], 5)
    add_figure(doc, FIGURES[5], 6)

    doc.add_page_break()
    doc.add_heading("Cross-dataset generalization", level=1)
    add_paragraph(
        doc,
        [
            ("Cervical WTA TopoLink-CCI full common-db run 生成 ", False, False),
            ("2,404,971 rows", True, False),
            ("。其 top hit 为 ", False, False),
            ("DSC2-DSG3 / Differentiating Tumor Cells -> Differentiating Tumor Cells", True, False),
            ("，与 Breast WTA 中 vascular/endothelial top axes 不同，支持 tissue-context-specific CCI prioritization。", False, False),
        ],
    )
    add_path_paragraph(
        doc,
        "Cervical artifact: ",
        "/cfs/klemming/scratch/h/hutaobo/pyxenium_cci_benchmark_2026-04/datasets/atera_cervical_wta/runs/full_common/pyxenium/pyxenium_standardized.tsv",
    )
    add_figure(doc, FIGURES[6], 7)
    add_figure(doc, FIGURES[7], 8)

    doc.add_page_break()
    doc.add_heading("Scope of Benchmark and Technical Limitations", level=1)
    add_paragraph(
        doc,
        [
            ("本 benchmark 不把所有外部 CCI 方法强行推进到 whole-dataset full run，而是为每个方法保留 full result、bounded subset result 或 reproducible failure card。SpaTalk smoke20k 完成 ", False, False),
            ("66/66 chunks", True, False),
            ("；NICHES pilot50k 完成 1,181,042 rows；Giotto pilot50k 完成 1,151,351 rows。这些结果说明真实方法链路可运行，同时也暴露了大规模 WTA 下的内存、sparse index 和 runtime scalability tradeoff。", False, False),
        ],
    )
    add_paragraph(
        doc,
        [
            ("SCILD 已在官方 source 被重新确认后通过 A100 专用 Python 3.11 环境升级为 bounded subset result：3000 cells × 20 CCI pairs，2,880 rows，峰值内存约 18.9GB。Copulacci 和 NicheNet 是 expanded 18-method benchmark 中仍保留的 reproducible failure cards；Copulacci 官方 source 可安装和 import，但没有可安全映射到 common-resource benchmark 的 native bounded workflow；NicheNet 在 PDC 和 A100 均受 R dependency/API audit 约束，并且其定位更接近 downstream receiver-response support，而不是直接 spatial CCI ranker。FastCCC 已通过 A100 retry 产生 20k bounded standardized output。", False, False),
        ],
    )
    add_figure(doc, FIGURES[8], 9)
    add_figure(doc, FIGURES[9], 10)

    doc.add_page_break()
    doc.add_heading("Methodological Discussion", level=1)
    add_paragraph(
        doc,
        [
            ("本 benchmark 不直接比较不同方法的 raw score，因为各方法的 score 定义不一致。评价体系采用 rank、canonical recovery、cross-method consistency、false-positive controls、Synthetic Truth 和 cross-dataset generalization。", False, False)
        ],
    )
    add_paragraph(
        doc,
        [
            ("F1 只用于 Synthetic Truth 或预定义 canonical truth set；ARI 不作为主指标，因为当前输出是 interaction-axis ranking，而不是统一 clustering assignment。", False, False)
        ],
    )
    add_paragraph(
        doc,
        [
            ("本研究建立的评估框架旨在为真实复杂组织切片中的 CCI 优先级排序提供可靠依据，而非单纯追求合成数据集上的指标拟合。", True, False)
        ],
    )
    add_paragraph(
        doc,
        [
            ("TopoLink-CCI 应被表述为 topology-guided spatial CCI hypothesis prioritization framework。所有结论均为 computational evidence，不应被解读为单靠 CCI_score 证明蛋白结合、因果通讯或功能效应。", False, False)
        ],
    )
    add_figure(doc, FIGURES[10], 11)
    add_figure(doc, FIGURES[11], 12)

    doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    build_report()
