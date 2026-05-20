"""Build a Nature-style short communication package for TopoLink-CCI.

This script follows the local ``nature-figure`` plotting contract:

- Python/Matplotlib only.
- Editable SVG text and Type 42 PDF fonts.
- One clear claim per panel.
- Source-data TSVs for every panel.
- Separate two-main-figure and one-main-figure-fallback outputs.
"""

from __future__ import annotations

import json
import math
import re
import shutil
import subprocess
from datetime import datetime, timezone
from pathlib import Path

import matplotlib

matplotlib.use("Agg")

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from matplotlib import patches
from matplotlib.lines import Line2D

try:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Inches, Pt
except Exception:  # pragma: no cover
    Document = None
    OxmlElement = None
    qn = None


REPO = Path(__file__).resolve().parents[4]
CCI_ROOT = REPO / "benchmarking" / "cci_2026_atera"
PACKAGE_ROOT = CCI_ROOT / "topolink_cci_short_communication" / "nature_short_communication_20260516"
FIG_DIR = PACKAGE_ROOT / "figures"
SOURCE_DIR = PACKAGE_ROOT / "source_data"
META_DIR = PACKAGE_ROOT / "metadata"
MANUSCRIPT_DIR = PACKAGE_ROOT / "manuscript"
QA_DIR = PACKAGE_ROOT / "qa"

METHOD_MATRIX = CCI_ROOT / "results" / "method_completion_matrix.tsv"
SYNTHETIC_TRUTH = (
    CCI_ROOT
    / "results"
    / "publication_benchmark_24h_20260511"
    / "pdc_collected"
    / "publication_benchmark_24h_20260511"
    / "synthetic_truth"
    / "tables"
    / "synthetic_truth_metrics.tsv"
)
EVIDENCE = CCI_ROOT / "pdc_validation_v2_collected" / "topolink_cci_validation_v2" / "tables" / "topolink_cci_validation_v2_evidence.tsv"
BREAST_TOP = CCI_ROOT / "results" / "preview_completed_20260511" / "source_data" / "fig3_breast_topolink_top_axes.tsv"
CROSS_DATASET = CCI_ROOT / "results" / "preview_completed_20260511" / "source_data" / "fig4_breast_cervical_topolink_comparison.tsv"
CANONICAL_RANK = CCI_ROOT / "results" / "cross_method_comparison_20260511" / "source_data" / "fig6_canonical_pair_rank_heatmap.tsv"
COMPONENTS = CCI_ROOT / "vwf_selp_deep_dive" / "tables" / "component_decomposition.tsv"
HOTSPOT_SUMMARY = CCI_ROOT / "vwf_selp_deep_dive" / "tables" / "hotspot_summary.tsv"
HOTSPOT_IMAGE = CCI_ROOT / "vwf_selp_deep_dive" / "figures" / "vwf_selp_hotspot_spatial_overlay.png"


COLORS = {
    "text": "#111827",
    "grid": "#E5E7EB",
    "muted": "#6B7280",
    "teal": "#0F766E",
    "blue": "#2563EB",
    "orange": "#D97706",
    "red": "#B91C1C",
    "green": "#15803D",
    "purple": "#6D28D9",
    "brown": "#8A5A2B",
    "gray": "#6B7280",
    "full": "#86EFAC",
    "bounded": "#FBBF24",
    "failure": "#FCA5A5",
}

THEME = {
    "VWF-SELP": "vascular",
    "VWF-LRP1": "vascular",
    "MMRN2-CD93": "vascular",
    "MMRN2-CLEC14A": "vascular",
    "COL4A2-CD93": "vascular",
    "EFNB2-PECAM1": "vascular",
    "VWF-ITGA9": "vascular",
    "HSPG2-LRP1": "stromal",
    "CD48-CD2": "immune",
    "DLL4-NOTCH3": "notch",
    "CXCL12-CXCR4": "stromal",
    "JAG1-NOTCH2": "tumor",
}

THEME_COLORS = {
    "vascular": COLORS["teal"],
    "stromal": COLORS["brown"],
    "immune": COLORS["green"],
    "notch": COLORS["purple"],
    "tumor": COLORS["gray"],
}


def configure_matplotlib() -> None:
    plt.rcParams.update(
        {
            "font.family": "sans-serif",
            "font.sans-serif": ["Arial", "Helvetica", "DejaVu Sans", "sans-serif"],
            "pdf.fonttype": 42,
            "ps.fonttype": 42,
            "svg.fonttype": "none",
            "font.size": 6.2,
            "axes.titlesize": 6.4,
            "axes.labelsize": 6.0,
            "xtick.labelsize": 5.2,
            "ytick.labelsize": 5.2,
            "legend.fontsize": 5.2,
            "axes.linewidth": 0.55,
            "axes.spines.top": False,
            "axes.spines.right": False,
            "legend.frameon": False,
            "savefig.dpi": 600,
        }
    )


def ensure_dirs() -> None:
    for path in [FIG_DIR, SOURCE_DIR, META_DIR, MANUSCRIPT_DIR, QA_DIR]:
        path.mkdir(parents=True, exist_ok=True)


def read_table(path: Path) -> pd.DataFrame:
    if not path.exists():
        raise FileNotFoundError(path)
    return pd.read_csv(path, sep="\t")


def load_data() -> dict[str, pd.DataFrame]:
    evidence = read_table(EVIDENCE)
    evidence["axis"] = evidence["ligand"] + "-" + evidence["receptor"]
    evidence["theme"] = evidence["axis"].map(THEME).fillna("tumor")
    data = {
        "method": read_table(METHOD_MATRIX),
        "synthetic": read_table(SYNTHETIC_TRUTH),
        "evidence": evidence,
        "breast_top": read_table(BREAST_TOP),
        "cross_dataset": read_table(CROSS_DATASET),
        "canonical_rank": read_table(CANONICAL_RANK),
        "components": read_table(COMPONENTS),
        "hotspot_summary": read_table(HOTSPOT_SUMMARY),
    }
    return data


def save_source(panel_id: str, df: pd.DataFrame) -> Path:
    path = SOURCE_DIR / f"{panel_id}.tsv"
    df.to_csv(path, sep="\t", index=False)
    return path


def save_figure(fig: plt.Figure, stem: str) -> dict[str, str]:
    outputs: dict[str, str] = {}
    for ext in ["pdf", "svg", "png", "tiff"]:
        path = FIG_DIR / f"{stem}.{ext}"
        if ext == "tiff":
            fig.savefig(path, dpi=600, bbox_inches="tight", pil_kwargs={"compression": "tiff_lzw"})
        else:
            fig.savefig(path, dpi=600, bbox_inches="tight")
        if ext == "svg":
            path.write_text("\n".join(line.rstrip() for line in path.read_text(encoding="utf-8").splitlines()) + "\n", encoding="utf-8")
        if ext in {"png", "tiff"}:
            try:
                from PIL import Image

                with Image.open(path) as image:
                    if image.mode != "RGB":
                        image.convert("RGB").save(path, dpi=image.info.get("dpi", (600, 600)))
            except Exception:
                pass
        outputs[ext] = str(path)
    return outputs


def panel_label(ax: plt.Axes, label: str) -> None:
    ax.text(
        -0.07,
        1.05,
        label,
        transform=ax.transAxes,
        fontsize=8,
        fontweight="bold",
        color=COLORS["text"],
        ha="left",
        va="top",
    )


def rounded_box(ax: plt.Axes, xy: tuple[float, float], wh: tuple[float, float], text: str, color: str) -> None:
    x, y = xy
    w, h = wh
    box = patches.FancyBboxPatch(
        (x, y),
        w,
        h,
        boxstyle="round,pad=0.02,rounding_size=0.02",
        fc="white",
        ec=color,
        lw=0.7,
    )
    ax.add_patch(box)
    ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=4.8, color=COLORS["text"])


def draw_problem_schematic(ax: plt.Axes) -> pd.DataFrame:
    panel_label(ax, "a")
    ax.set_title("False-positive risks", loc="left", fontweight="bold")
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    boxes = [
        ("co-expression\nfalse positives", 0.06, 0.66, COLORS["red"]),
        ("cell abundance\nbias", 0.58, 0.66, COLORS["orange"]),
        ("proximity without\nmolecular support", 0.06, 0.28, COLORS["blue"]),
        ("CCI resource\nambiguity", 0.58, 0.28, COLORS["purple"]),
    ]
    for text, x, y, color in boxes:
        rounded_box(ax, (x, y), (0.36, 0.16), text, color)
    ax.text(0.50, 0.50, "False-positive\nrisk", ha="center", va="center", fontsize=7, fontweight="bold", color=COLORS["text"])
    for _, x, y, color in boxes:
        ax.annotate("", xy=(0.50, 0.50), xytext=(x + 0.18, y + 0.08), arrowprops=dict(arrowstyle="->", lw=0.7, color=color))
    return pd.DataFrame(boxes, columns=["risk", "x", "y", "color"])


def draw_score_architecture(ax: plt.Axes) -> pd.DataFrame:
    panel_label(ax, "b")
    ax.set_title("Score architecture", loc="left", fontweight="bold")
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    components = [
        ("sender\nanchor", 0.08, 0.68, COLORS["teal"]),
        ("receiver\nanchor", 0.38, 0.68, COLORS["purple"]),
        ("structure\nbridge", 0.68, 0.68, COLORS["blue"]),
        ("sender\nexpression", 0.08, 0.36, COLORS["green"]),
        ("receiver\nexpression", 0.38, 0.36, COLORS["green"]),
        ("local\ncontact", 0.68, 0.36, COLORS["orange"]),
    ]
    for text, x, y, color in components:
        rounded_box(ax, (x, y), (0.22, 0.14), text, color)
        ax.annotate("", xy=(0.50, 0.20), xytext=(x + 0.11, y), arrowprops=dict(arrowstyle="->", lw=0.55, color=color))
    ax.text(
        0.50,
        0.13,
        "prior-weighted geometric mean\n= discovery score, not proof",
        ha="center",
        va="center",
        fontsize=6,
        bbox=dict(fc="#F9FAFB", ec="#D1D5DB", boxstyle="round,pad=0.25"),
    )
    return pd.DataFrame(components, columns=["component", "x", "y", "color"])


def draw_workflow(ax: plt.Axes) -> pd.DataFrame:
    panel_label(ax, "c")
    ax.set_title("Workflow", loc="left", fontweight="bold")
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    steps = [
        ("Atera\nWTA", 0.04, 0.42),
        ("topology\nmap", 0.24, 0.42),
        ("CCI\nranking", 0.44, 0.42),
        ("controls", 0.64, 0.42),
        ("evidence\nclass", 0.82, 0.42),
    ]
    for i, (text, x, y) in enumerate(steps):
        rounded_box(ax, (x, y), (0.15, 0.18), text, COLORS["teal"] if i < 3 else COLORS["orange"])
        if i < len(steps) - 1:
            ax.annotate("", xy=(steps[i + 1][1] - 0.01, y + 0.09), xytext=(x + 0.16, y + 0.09), arrowprops=dict(arrowstyle="->", lw=0.7, color=COLORS["text"]))
    ax.text(0.50, 0.19, "Breast WTA: 170,057 cells; 3,299 common CCI pairs; 1,319,600 hypotheses", ha="center", fontsize=5.8)
    return pd.DataFrame(steps, columns=["step", "x", "y"])


def draw_synthetic(ax: plt.Axes, synthetic: pd.DataFrame) -> pd.DataFrame:
    panel_label(ax, "d")
    plot = synthetic.copy()
    label_map = {
        "TopoLink-CCI": "TopoLink\nCCI",
        "expression_only": "expr.",
        "contact_only": "contact",
        "topology_anchor_only": "anchor",
    }
    plot["display"] = plot["method"].map(label_map).fillna(plot["method"])
    x = np.arange(len(plot))
    width = 0.34
    ax.bar(x - width / 2, plot["auroc"], width=width, color=COLORS["blue"], label="AUROC")
    ax.bar(x + width / 2, plot["auprc"], width=width, color=COLORS["teal"], label="AUPRC")
    ax.set_xticks(x)
    ax.set_xticklabels(plot["display"])
    ax.set_ylim(0.45, 1.03)
    ax.set_ylabel("metric")
    ax.set_title("Synthetic Truth", loc="left", fontweight="bold")
    ax.grid(axis="y", color=COLORS["grid"], lw=0.5)
    ax.legend(loc="lower left")
    top = plot.loc[plot["method"].eq("TopoLink-CCI")].iloc[0]
    anchor = plot.loc[plot["method"].eq("topology_anchor_only")].iloc[0]
    ax.text(
        0.03,
        0.95,
        f"TopoLink-CCI\nAUROC {top.auroc:.4f}\nAUPRC {top.auprc:.4f}",
        transform=ax.transAxes,
        ha="left",
        va="top",
        fontsize=5.4,
        bbox=dict(fc="white", ec="#D1D5DB", boxstyle="round,pad=0.25"),
    )
    ax.text(0.98, 0.52, f"anchor-only\nAUPRC {anchor.auprc:.4f}", transform=ax.transAxes, ha="right", va="bottom", fontsize=5.4, color=COLORS["orange"])
    return plot[["method", "auroc", "auprc", "precision_at_5", "recall_at_5", "f1_at_5"]]


def draw_evidence_matrix(ax: plt.Axes, evidence: pd.DataFrame) -> pd.DataFrame:
    panel_label(ax, "e")
    layers = [
        ("expr", "expression_specificity_support"),
        ("label", "cell_label_permutation_support"),
        ("spatial", "spatial_null_support"),
        ("matched", "matched_gene_control_support"),
        ("target", "downstream_target_support"),
        ("signal", "functional_received_signal_support"),
        ("cross", "cross_method_support"),
        ("boot", "bootstrap_stability_support"),
    ]
    df = evidence.sort_values("pyxenium_rank").head(7).copy()
    matrix = df[[col for _, col in layers]].astype(bool).to_numpy()
    ax.set_title("Validation controls", loc="left", fontweight="bold")
    ax.set_xlim(-0.5, len(layers) - 0.5)
    ax.set_ylim(len(df) - 0.5, -0.5)
    ax.set_xticks(range(len(layers)))
    ax.set_xticklabels([name for name, _ in layers], rotation=45, ha="right")
    ax.set_yticks(range(len(df)))
    ax.set_yticklabels(df["axis"].tolist())
    for i in range(matrix.shape[0]):
        for j in range(matrix.shape[1]):
            supported = matrix[i, j]
            ax.scatter(
                j,
                i,
                s=52,
                marker="o" if supported else "X",
                color=COLORS["green"] if supported else COLORS["orange"],
                edgecolor="white",
                linewidth=0.45,
            )
    ax.grid(color=COLORS["grid"], lw=0.45)
    ax.tick_params(length=0)
    ax.text(0.99, 0.02, "7/7 strong; 0 artifact risk", transform=ax.transAxes, ha="right", va="bottom", fontsize=6, fontweight="bold")
    keep = ["axis", "pyxenium_rank", "CCI_score", "support_count", "evidence_class"] + [col for _, col in layers]
    return df[keep]


def draw_breast_top_axes(ax: plt.Axes, breast_top: pd.DataFrame, label: str = "a") -> pd.DataFrame:
    panel_label(ax, label)
    interpretable_axes = [
        "VWF-SELP",
        "VWF-LRP1",
        "EFNB2-PECAM1",
        "MMRN2-CLEC14A",
        "HSPG2-LRP1",
        "COL4A2-CD93",
        "MMRN2-CD93",
        "VWF-ITGA9",
    ]
    plot = breast_top.loc[breast_top["axis"].isin(interpretable_axes)].copy()
    if plot.empty:
        plot = breast_top.copy().sort_values("CCI_score").tail(8)
    plot = plot.sort_values("CCI_score")
    plot["theme"] = plot["axis"].map(THEME).fillna("tumor")
    colors = [THEME_COLORS[t] for t in plot["theme"]]
    ax.barh(plot["axis"], plot["CCI_score"], color=colors)
    ax.set_xlim(0.63, 0.82)
    ax.set_xlabel("CCI score")
    ax.set_title("Breast top axes", loc="left", fontweight="bold")
    ax.grid(axis="x", color=COLORS["grid"], lw=0.5)
    return plot[["dataset", "axis", "sender_receiver", "CCI_score", "local_contact", "cross_edge_count", "theme"]]


def draw_vwf_decomposition(ax: plt.Axes, components: pd.DataFrame) -> pd.DataFrame:
    panel_label(ax, "b")
    order = ["sender_anchor", "receiver_anchor", "structure_bridge", "sender_expr", "receiver_expr", "local_contact"]
    df = components.set_index("component").loc[order].reset_index()
    colors = [COLORS["teal"], COLORS["purple"], COLORS["blue"], COLORS["green"], COLORS["green"], COLORS["orange"]]
    ax.barh(df["component"].str.replace("_", "\n"), df["value"], color=colors)
    ax.set_xlim(0, 1.05)
    ax.set_xlabel("component value")
    ax.set_title("VWF-SELP score components", loc="left", fontweight="bold")
    ax.grid(axis="x", color=COLORS["grid"], lw=0.5)
    ax.text(
        0.05,
        0.10,
        "CCI_score 0.791\nlocal_contact 0.291\n12,779 endothelial-endothelial edges",
        transform=ax.transAxes,
        ha="left",
        va="bottom",
        fontsize=5.5,
        bbox=dict(fc="white", ec="#D1D5DB", boxstyle="round,pad=0.25"),
    )
    return df


def draw_hotspot(ax: plt.Axes, hotspot_summary: pd.DataFrame) -> pd.DataFrame:
    panel_label(ax, "c")
    ax.set_title("Spatial hotspots", loc="left", fontweight="bold")
    image = plt.imread(HOTSPOT_IMAGE)
    ax.imshow(image, rasterized=True)
    ax.axis("off")
    n_hot = int(hotspot_summary.loc[0, "n_hotspot_endothelial_cells"])
    ax.text(
        0.03,
        0.96,
        f"{n_hot} hotspot endothelial cells\nRNA-level spatial evidence",
        transform=ax.transAxes,
        ha="left",
        va="top",
        fontsize=5.4,
        bbox=dict(fc="white", ec="#D1D5DB", boxstyle="round,pad=0.25"),
    )
    return hotspot_summary.assign(image_path=str(HOTSPOT_IMAGE))


def draw_method_status(ax: plt.Axes, method: pd.DataFrame, label: str = "d") -> pd.DataFrame:
    panel_label(ax, label)
    breast = method.loc[method["dataset"].eq("atera_breast_wta")].copy()
    order = ["full_result", "bounded_subset_result", "reproducible_failure_card"]
    counts = breast["status"].value_counts().reindex(order).fillna(0).astype(int)
    labels = ["full", "bounded", "failure\ncard"]
    colors = [COLORS["full"], COLORS["bounded"], COLORS["failure"]]
    ax.bar(labels, counts.values, color=colors, edgecolor="#374151", linewidth=0.5)
    ax.set_ylim(0, max(counts.values) + 2)
    ax.set_ylabel("methods")
    ax.set_title("Benchmark status", loc="left", fontweight="bold")
    ax.grid(axis="y", color=COLORS["grid"], lw=0.5)
    for i, value in enumerate(counts.values):
        ax.text(i, value + 0.25, str(value), ha="center", va="bottom", fontsize=7, fontweight="bold")
    ax.text(
        0.98,
        0.96,
        "18 Breast methods\n9 full + 9 bounded\n0 failure / 0 deferred",
        transform=ax.transAxes,
        ha="right",
        va="top",
        fontsize=6,
        fontweight="bold",
    )
    return counts.rename_axis("status").reset_index(name="n_methods")


def draw_canonical_rank(ax: plt.Axes, canonical: pd.DataFrame, label: str = "e") -> pd.DataFrame:
    panel_label(ax, label)
    selected_methods = ["TopoLink-CCI", "CellPhoneDB", "LARIS", "LIANA+", "SpatialDM", "stLearn", "Squidpy"]
    canonical_ids = ["VWF|SELP", "VWF|LRP1", "MMRN2|CD93", "DLL4|NOTCH3", "CXCL12|CXCR4"]
    df = canonical.loc[canonical["method"].isin(selected_methods) & canonical["canonical_id"].isin(canonical_ids)].copy()
    pivot = df.pivot_table(index="method", columns="canonical_id", values="rank_capped", aggfunc="min").reindex(selected_methods)
    pivot = pivot.reindex(columns=canonical_ids)
    ax.set_xlim(-0.5, len(canonical_ids) - 0.5)
    ax.set_ylim(len(selected_methods) - 0.5, -0.5)
    ax.set_xticks(range(len(pivot.columns)))
    ax.set_xticklabels([c.replace("|", "-").replace("-", "\n") for c in pivot.columns])
    ax.set_yticks(range(len(pivot.index)))
    ax.set_yticklabels(pivot.index)
    ax.set_title("Canonical ranks", loc="left", fontweight="bold")
    ax.grid(True, color=COLORS["grid"], lw=0.45)
    for i, method in enumerate(pivot.index):
        for j, axis_id in enumerate(pivot.columns):
            value = pivot.loc[method, axis_id]
            if pd.isna(value):
                ax.plot(j, i, marker="x", color="#BDBDBD", markersize=4.0, mew=0.8)
                continue
            if value <= 50:
                color, size = COLORS["teal"], 58
            elif value <= 500:
                color, size = "#4D96A8", 42
            elif value <= 5000:
                color, size = COLORS["orange"], 26
            else:
                color, size = "#C9C9C9", 14
            ax.scatter(j, i, s=size, color=color, edgecolor="white", linewidth=0.35, zorder=3)
            ax.text(j, i + 0.29, f"{int(value)}", ha="center", va="center", fontsize=3.8, color=COLORS["text"])
    handles = [
        Line2D([0], [0], marker="o", color="none", markerfacecolor=COLORS["teal"], markeredgecolor="white", markersize=5, label="top 50"),
        Line2D([0], [0], marker="o", color="none", markerfacecolor="#4D96A8", markeredgecolor="white", markersize=4, label="top 500"),
        Line2D([0], [0], marker="o", color="none", markerfacecolor=COLORS["orange"], markeredgecolor="white", markersize=3, label="top 5000"),
        Line2D([0], [0], marker="x", color="#BDBDBD", markersize=4, label="not detected"),
    ]
    ax.legend(handles=handles, loc="upper left", bbox_to_anchor=(1.02, 1.0), borderaxespad=0.0)
    return df


def draw_cross_dataset(ax: plt.Axes, cross: pd.DataFrame, label: str = "f") -> pd.DataFrame:
    panel_label(ax, label)
    top = cross.groupby("dataset", group_keys=False).head(3).copy()
    top = top.sort_values(["dataset", "CCI_score"], ascending=[True, True])
    top["display"] = top["dataset"].str.title() + ": " + top["axis"]
    colors = [COLORS["teal"] if d == "breast" else COLORS["purple"] for d in top["dataset"]]
    ax.barh(top["display"], top["CCI_score"], color=colors)
    ax.set_xlim(0.60, max(0.84, float(top["CCI_score"].max()) + 0.02))
    ax.set_xlabel("CCI score")
    ax.set_title("Breast vs cervical", loc="left", fontweight="bold")
    ax.grid(axis="x", color=COLORS["grid"], lw=0.5)
    ax.text(
        0.02,
        0.03,
        "Breast top: VWF-SELP endothelial activation\nCervical top: DSC2-DSG3 differentiating tumor adhesion\nCervical full: 2,404,971 rows",
        transform=ax.transAxes,
        ha="left",
        va="bottom",
        fontsize=5.3,
        bbox=dict(fc="white", ec="#D1D5DB", boxstyle="round,pad=0.25"),
    )
    return top[["dataset", "axis", "sender_receiver", "CCI_score", "local_contact", "cross_edge_count"]]


def draw_bounded_scalability(ax: plt.Axes, method: pd.DataFrame, label: str = "e") -> pd.DataFrame:
    panel_label(ax, label)
    order = ["Giotto", "SpaTalk", "NICHES", "CellNEST", "CellAgentChat", "FastCCC", "SCILD", "Copulacci", "NicheNet"]
    df = method.loc[
        method["dataset"].eq("atera_breast_wta")
        & method["status"].eq("bounded_subset_result")
        & method["method"].isin(order)
    ].copy()
    df["method"] = pd.Categorical(df["method"], categories=order, ordered=True)
    df = df.sort_values("method")
    df["n_rows_numeric"] = pd.to_numeric(df["n_rows"], errors="coerce")
    df["log10_rows"] = np.log10(df["n_rows_numeric"].clip(lower=1))
    colors = []
    for method_name in df["method"].astype(str):
        if method_name in {"NicheNet"}:
            colors.append(COLORS["purple"])
        elif method_name in {"FastCCC", "SCILD", "Copulacci"}:
            colors.append(COLORS["blue"])
        else:
            colors.append(COLORS["orange"])
    ax.barh(df["method"].astype(str), df["log10_rows"], color=colors, edgecolor="#374151", linewidth=0.35)
    ax.invert_yaxis()
    ax.set_xlabel("log10 standardized rows")
    ax.set_title("Bounded appendix scalability", loc="left", fontweight="bold")
    ax.grid(axis="x", color=COLORS["grid"], lw=0.5)
    for i, row in enumerate(df.itertuples()):
        value = getattr(row, "n_rows_numeric")
        label_text = f"{int(value):,}" if pd.notna(value) else "reported"
        ax.text(getattr(row, "log10_rows") + 0.04, i, f"{label_text} ({row.phase})", va="center", fontsize=4.8, color=COLORS["text"])
    ax.text(
        0.02,
        0.03,
        "Bounded evidence is terminal appendix evidence,\nnot a full whole-dataset equivalence claim.",
        transform=ax.transAxes,
        ha="left",
        va="bottom",
        fontsize=5.2,
        bbox=dict(fc="white", ec="#D1D5DB", boxstyle="round,pad=0.25"),
    )
    return df[["dataset", "method", "status", "evidence_level", "phase", "n_rows", "remote_or_local_path", "notes", "log10_rows"]]


def make_figure_1(data: dict[str, pd.DataFrame]) -> dict[str, str]:
    fig = plt.figure(figsize=(7.15, 5.35), constrained_layout=False)
    gs = fig.add_gridspec(2, 3, width_ratios=[1.0, 1.0, 1.15], height_ratios=[1.0, 1.2])
    fig.subplots_adjust(left=0.07, right=0.985, bottom=0.09, top=0.91, wspace=0.50, hspace=0.42)
    sources = {
        "fig1a_problem_schematic": save_source("fig1a_problem_schematic", draw_problem_schematic(fig.add_subplot(gs[0, 0]))),
        "fig1b_score_architecture": save_source("fig1b_score_architecture", draw_score_architecture(fig.add_subplot(gs[0, 1]))),
        "fig1c_workflow": save_source("fig1c_workflow", draw_workflow(fig.add_subplot(gs[0, 2]))),
        "fig1d_synthetic_truth": save_source("fig1d_synthetic_truth", draw_synthetic(fig.add_subplot(gs[1, 0]), data["synthetic"])),
        "fig1e_evidence_matrix": save_source("fig1e_evidence_matrix", draw_evidence_matrix(fig.add_subplot(gs[1, 1:]), data["evidence"])),
    }
    fig.text(0.07, 0.985, "Figure 1 | TopoLink-CCI method and validation logic", fontsize=7.2, fontweight="bold", va="top")
    outputs = save_figure(fig, "figure_1_topolink_cci_method_validation")
    plt.close(fig)
    write_manifest("figure_1", outputs, sources, "TopoLink-CCI needs topology-expression-contact concordance plus independent controls.")
    return outputs


def make_figure_2(data: dict[str, pd.DataFrame]) -> dict[str, str]:
    fig = plt.figure(figsize=(7.15, 6.65), constrained_layout=False)
    gs = fig.add_gridspec(3, 2, width_ratios=[0.95, 1.1], height_ratios=[0.88, 1.12, 1.0])
    fig.subplots_adjust(left=0.10, right=0.96, bottom=0.075, top=0.94, wspace=0.66, hspace=0.78)
    sources = {
        "fig2a_method_status": save_source("fig2a_method_status", draw_method_status(fig.add_subplot(gs[0, 0]), data["method"], label="a")),
        "fig2b_breast_top_axes": save_source("fig2b_breast_top_axes", draw_breast_top_axes(fig.add_subplot(gs[0, 1]), data["breast_top"], label="b")),
        "fig2c_canonical_rank": save_source("fig2c_canonical_rank", draw_canonical_rank(fig.add_subplot(gs[1, 0]), data["canonical_rank"], label="c")),
        "fig2d_cross_dataset": save_source("fig2d_cross_dataset", draw_cross_dataset(fig.add_subplot(gs[1, 1]), data["cross_dataset"], label="d")),
        "fig2e_bounded_scalability": save_source("fig2e_bounded_scalability", draw_bounded_scalability(fig.add_subplot(gs[2, :]), data["method"], label="e")),
    }
    fig.text(0.10, 0.987, "Figure 2 | Whole-dataset benchmarking and biological interpretation", fontsize=7.2, fontweight="bold", va="top")
    outputs = save_figure(fig, "figure_2_topolink_cci_discovery_benchmark")
    plt.close(fig)
    write_manifest("figure_2", outputs, sources, "Whole-dataset TopoLink-CCI outputs recover interpretable and tissue-specific CCI axes.")
    return outputs


def make_fallback(data: dict[str, pd.DataFrame]) -> dict[str, str]:
    fig = plt.figure(figsize=(7.15, 6.65), constrained_layout=False)
    gs = fig.add_gridspec(3, 2, height_ratios=[0.9, 1.2, 1.05], width_ratios=[1.0, 1.15])
    fig.subplots_adjust(left=0.08, right=0.97, bottom=0.07, top=0.94, wspace=0.52, hspace=0.72)
    sources = {
        "fallback_a_score": save_source("fallback_a_score_architecture", draw_score_architecture(fig.add_subplot(gs[0, 0]))),
        "fallback_b_synthetic": save_source("fallback_b_synthetic_truth", draw_synthetic(fig.add_subplot(gs[0, 1]), data["synthetic"])),
        "fallback_c_evidence": save_source("fallback_c_evidence_matrix", draw_evidence_matrix(fig.add_subplot(gs[1, :]), data["evidence"])),
        "fallback_d_status": save_source("fallback_d_method_status", draw_method_status(fig.add_subplot(gs[2, 0]), data["method"], label="d")),
        "fallback_e_cross_dataset": save_source("fallback_e_cross_dataset", draw_cross_dataset(fig.add_subplot(gs[2, 1]), data["cross_dataset"], label="e")),
    }
    fig.text(0.08, 0.987, "TopoLink-CCI prioritizes topology-supported spatial CCI hypotheses", fontsize=7.2, fontweight="bold", va="top")
    outputs = save_figure(fig, "figure_1_fallback_single_main")
    plt.close(fig)
    write_manifest("figure_1_fallback", outputs, sources, "Single-display-item fallback combining method, validation and discovery evidence.")
    return outputs


def write_manifest(figure_id: str, outputs: dict[str, str], sources: dict[str, Path], claim: str) -> None:
    manifest = {
        "figure_id": figure_id,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "claim": claim,
        "outputs": outputs,
        "source_data": {key: str(value) for key, value in sources.items()},
        "figure_contract": {
            "backend": "Python/Matplotlib",
            "pdf_fonttype": 42,
            "svg_fonttype": "none",
            "source_data_required": True,
            "journal_target": "Nature Methods Brief Communication",
        },
    }
    (META_DIR / f"{figure_id}_manifest.json").write_text(json.dumps(manifest, indent=2) + "\n", encoding="utf-8")


def write_manuscripts() -> dict[str, Path]:
    abstract = (
        "Spatial cell-cell interaction inference is confounded by co-expression, cell abundance and incomplete molecular resources. "
        "We introduce TopoLink-CCI, a topology-guided framework that integrates tissue topology, expression specificity and local contact, "
        "then challenges candidates with orthogonal controls. In public 10x Genomics Atera Whole Transcriptome Assay (Atera WTA) FFPE breast cancer data, an expanded 18-method benchmark reached nine full "
        "and nine bounded terminal results with no failure or deferred methods; Atera WTA cervical cancer analysis showed tissue-context-specific tumor adhesion axes."
    )
    main_text = f"""# Topology-guided prioritization of spatial cell-cell interaction axes

**Target journal:** Nature Methods Brief Communication
**Fallback journal:** Nature Biotechnology Brief Communication
**Method:** TopoLink-CCI

## Abstract

{abstract}

## Main text

Spatial transcriptomics has made it possible to infer cell-cell interaction (CCI) axes in intact tissues, but high-resolution whole-transcriptome imaging also increases false-positive risk. Ligand and receptor genes can appear biologically plausible because both are strongly expressed, because the corresponding cell types are abundant, or because a curated molecular resource includes extracellular-matrix, adhesion, scavenger-receptor and shared activation-state relationships that are not classical secreted signaling events. Spatial proximity improves interpretation but does not by itself prove molecular exchange, receptor activation or causal signaling. The central problem is therefore not only to rank candidate molecular pairs, but to prioritize CCI hypotheses that are compatible with tissue topology, local organization and independent controls.

We developed TopoLink-CCI as a topology-guided CCI prioritization framework in pyXenium. In CCI-resource mode, each ligand, receptor, sender and receiver combination is scored by six retained components: sender topology anchor, receiver topology anchor, sender-receiver structure bridge, sender expression support, receiver expression support and local contact support. The discovery score is a prior-weighted geometric mean of these components, so a high score requires concordance across topology, expression and spatial contact rather than any single high-expression feature. Because every component is exported, users can diagnose whether an axis is topology-driven, expression-driven, contact-sensitive or potentially dominated by a resource prior. This design explicitly treats the score as a discovery statistic, not as proof of protein-level communication.

We first evaluated the scoring logic in a topology-preserving Synthetic Truth benchmark. The full TopoLink-CCI model achieved AUROC 0.9919 and AUPRC 0.8333, whereas the topology-anchor-only model retained high AUROC (0.9839) but dropped to AUPRC 0.5833. We then applied TopoLink-CCI to the public 10x Genomics Preview Data: Atera In Situ Gene Expression, FFPE Human Breast Cancer dataset, generated using the pre-commercial Atera Whole Transcriptome Assay (Atera WTA) and Atera Onboard Analysis development workflow. TopoLink-CCI generated 1,319,600 common-resource CCI hypotheses from 170,057 cells in this breast cancer tissue. The expanded breast benchmark now includes 18 terminalized methods: nine full whole-dataset results and nine bounded subset results, with zero failure cards and zero deferred methods. Late-stage rescue runs converted FastCCC, SCILD, Copulacci and NicheNet into bounded appendix evidence, with NicheNet interpreted specifically as downstream receiver-response support rather than a direct spatial CCI ranker. Seven representative axes were evaluated with orthogonal controls inspired by established CCI methods, including cell-label permutation, spatial nulls, matched-expression gene controls, downstream target support, received-signal association, cross-method consensus, component ablation and bootstrap stability.

TopoLink-CCI prioritized biologically interpretable axes spanning vascular activation, stromal matrix biology, immune adhesion, Notch signaling and tumor-intrinsic adhesion. The top breast axis, VWF-SELP from endothelial cells to endothelial cells, is best interpreted as an endothelial activation and vascular adhesion niche consistent with Weibel-Palade body biology, not as direct proof of VWF/P-selectin protein release. Cross-dataset application to the corresponding public 10x Genomics Preview Data: Atera In Situ Gene Expression, FFPE Human Cervical Cancer dataset produced 2,404,971 hypotheses from 717,576 cells and a distinct top tumor-adhesion axis, DSC2-DSG3 in differentiating tumor cells, supporting tissue-context-specific prioritization. TopoLink-CCI is therefore a spatial CCI hypothesis-prioritization framework: it narrows a large molecular search space to axes consistent with tissue topology, expression specificity, local contact and independent computational controls, while leaving causal signaling and protein-level mechanism to orthogonal experimental validation.

## Figure legends

**Figure 1 | TopoLink-CCI method and validation logic.** **a,** Spatial CCI inference is vulnerable to co-expression, cell-abundance, proximity and resource-composition false positives. **b,** TopoLink-CCI scores each candidate using sender and receiver topology anchors, structure bridging, expression support and local contact. **c,** Workflow from Atera WTA preview datasets and pyXenium topology maps to ranked CCI axes and validation controls. **d,** Synthetic Truth evaluation shows that the full model achieves AUROC 0.9919 and AUPRC 0.8333, whereas topology-anchor-only scoring loses precision-recall performance. **e,** Orthogonal evidence matrix for seven interpretable breast cancer axes.

**Figure 2 | Whole-dataset benchmarking and biological interpretation.** **a,** The expanded Breast WTA benchmark terminalizes 18 methods as nine full whole-dataset results and nine bounded subset results, with zero failure cards and zero deferred methods. **b,** Top interpretable Breast WTA axes ranked by TopoLink-CCI score are led by VWF-SELP. **c,** Canonical recovery is compared by within-method rank rather than raw score. **d,** Breast and cervical WTA datasets yield tissue-context-specific top axes, including DSC2-DSG3 in cervical differentiating tumor cells. **e,** Bounded appendix methods provide scalability-aware terminal evidence, including FastCCC, SCILD, Copulacci and NicheNet late-stage rescue results.

## Guardrails

- TopoLink-CCI score is a discovery score, not proof of ligand binding, protein secretion, receptor activation or causal signaling.
- Raw scores are not compared across methods.
- Full whole-dataset and bounded subset methods are explicitly separated.
- F1, AUROC and AUPRC are used only for Synthetic Truth or predefined canonical truth sets.
"""
    methods = """# Online Methods outline

## TopoLink-CCI score

TopoLink-CCI evaluates each candidate molecular interaction axis using six retained components: sender topology anchor, receiver topology anchor, structure bridge, sender expression, receiver expression and local contact. The discovery score is computed as a prior-weighted geometric mean of these components.

## Topology map and expression support

Topology anchors are derived from pyXenium topology outputs linking genes, cell groups and tissue structures. Expression support is evaluated within sender and receiver cell groups and exported as diagnostic fields.

## Local contact graph

Local contact support is computed on a spatial neighbor graph and records active-edge support between sender and receiver populations.

## Synthetic Truth benchmark

Synthetic Truth data preserve tissue topology while implanting known CCI axes. AUROC, AUPRC and top-k precision/recall are used only because positive and negative axes are defined by construction.

## Benchmark status tiers

Methods were terminalized as full result or bounded subset result in the expanded Breast WTA benchmark. Reproducible failure cards remain the predefined stopping rule, but no expanded Breast method currently remains in that class. Bounded methods are not treated as equivalent to whole-dataset full methods. NicheNet is analyzed as downstream receiver-response support and is not presented as a direct spatial CCI ranker.
"""
    inquiry = """# Presubmission inquiry draft

Dear Editors,

We would like to ask whether you would consider a Brief Communication describing TopoLink-CCI, a topology-guided framework for spatial cell-cell interaction inference in whole-transcriptome imaging data.

TopoLink-CCI addresses a central limitation of current CCI analysis: co-expression, cell abundance and spatial proximity can generate plausible but weakly controlled interaction hypotheses. The method integrates tissue topology, expression specificity and local contact into an interpretable discovery score, then evaluates candidates with orthogonal false-positive controls.

In the public 10x Genomics Preview Data: Atera In Situ Gene Expression, FFPE Human Breast Cancer dataset generated using Atera Whole Transcriptome Assay (Atera WTA), TopoLink-CCI generated 1,319,600 CCI hypotheses and prioritized vascular, stromal, immune and Notch axes with strong computational support. In a Synthetic Truth benchmark it achieved AUROC 0.9919 and AUPRC 0.8333, outperforming topology-anchor-only scoring in precision-recall ranking. Cross-dataset application to the corresponding Atera WTA FFPE Human Cervical Cancer preview dataset produced 2,404,971 hypotheses and a distinct top tumor-adhesion axis, supporting tissue-context-specific prioritization.

The accompanying expanded benchmark terminalizes all 18 Breast WTA methods, comprising nine full whole-dataset results and nine bounded subset results with no remaining failure or deferred methods. Bounded results are reported as scalability-aware appendix evidence and remain explicitly separated from full whole-dataset comparisons.

We believe this work fits Nature Methods because it presents a concise method and validation framework for spatial omics, with broad relevance to tissue-scale CCI analysis. The proposed submission would include two main figures, online methods, source data and a complete benchmark status table.

Sincerely,
"""
    paths = {
        "main": MANUSCRIPT_DIR / "topolink_cci_short_communication.md",
        "methods": MANUSCRIPT_DIR / "online_methods.md",
        "inquiry": MANUSCRIPT_DIR / "presubmission_inquiry.md",
    }
    paths["main"].write_text(main_text, encoding="utf-8", newline="\n")
    paths["methods"].write_text(methods, encoding="utf-8", newline="\n")
    paths["inquiry"].write_text(inquiry, encoding="utf-8", newline="\n")
    return paths


def write_supplement_manifest(data: dict[str, pd.DataFrame]) -> Path:
    rows = [
        ("Supplementary Table 1", "method_completion_matrix", str(METHOD_MATRIX), "Final full/bounded terminal status for all benchmarked methods."),
        ("Supplementary Table 2", "validation_evidence", str(EVIDENCE), "False-positive controls and evidence classes for seven CCI axes."),
        ("Supplementary Table 3", "synthetic_truth_metrics", str(SYNTHETIC_TRUTH), "Synthetic Truth AUROC/AUPRC and top-k metrics."),
        ("Supplementary Table 4", "breast_cervical_top_axes", str(CROSS_DATASET), "Cross-dataset TopoLink-CCI top axes."),
        ("Supplementary source image", "vwf_selp_hotspot_overlay", str(HOTSPOT_IMAGE), "Spatial hotspot image retained for VWF-SELP deep-dive support."),
    ]
    manifest = pd.DataFrame(rows, columns=["item", "name", "path", "description"])
    path = MANUSCRIPT_DIR / "supplementary_data_manifest.tsv"
    manifest.to_csv(path, sep="\t", index=False)
    return path


def word_count(text: str) -> int:
    return len(re.findall(r"[A-Za-z0-9]+(?:[-'][A-Za-z0-9]+)?", text))


def add_line_numbering(section) -> None:
    if OxmlElement is None:  # pragma: no cover
        return
    sect_pr = section._sectPr
    existing = sect_pr.find(qn("w:lnNumType"))
    if existing is None:
        existing = OxmlElement("w:lnNumType")
        sect_pr.append(existing)
    existing.set(qn("w:countBy"), "1")
    existing.set(qn("w:start"), "1")
    existing.set(qn("w:restart"), "continuous")


def add_page_number(section) -> None:
    if OxmlElement is None:  # pragma: no cover
        return
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run("Page ")
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])


def set_submission_styles(doc: Document) -> None:
    for style_name, size in [("Normal", 12), ("Title", 16), ("Heading 1", 12), ("Heading 2", 12)]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
    doc.styles["Title"].font.bold = True
    doc.styles["Heading 1"].font.bold = True
    doc.styles["Heading 2"].font.bold = True


def add_double_spaced_paragraph(doc: Document, text: str, style: str | None = None, bold: bool = False) -> None:
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.line_spacing = 2.0
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold


def submission_text_sections() -> dict[str, str | list[str]]:
    abstract = (
        "Spatial cell-cell interaction inference is confounded by co-expression, cell abundance and incomplete molecular resources. "
        "TopoLink-CCI integrates tissue topology, expression specificity and local contact, then challenges candidates with orthogonal controls. "
        "Across public 10x Genomics Atera Whole Transcriptome Assay (Atera WTA) FFPE breast and cervical cancer preview datasets, it scales to whole tissues and prioritizes computationally supported vascular, stromal, immune and tumor interaction hypotheses."
    )
    body = [
        (
            "Spatial transcriptomics enables cell-cell interaction (CCI) inference in intact tissues, but whole-transcriptome imaging also expands false-positive risk. "
            "Candidate axes may rank highly because both genes are abundant, because cell groups are frequent, or because curated molecular resources include extracellular-matrix, adhesion, scavenger-receptor and shared activation-state relationships that are not classical secreted signaling. "
            "Spatial proximity improves plausibility but does not prove molecular exchange, receptor activation or causal communication."
        ),
        (
            "TopoLink-CCI is a topology-guided CCI prioritization framework in pyXenium. "
            "In CCI-resource mode, each ligand, receptor, sender and receiver combination is scored using sender topology anchor, receiver topology anchor, structure bridge, sender expression, receiver expression and local contact. "
            "The prior-weighted geometric mean favors concordance across topology, expression and spatial contact, and every component is exported so users can diagnose topology-driven, expression-driven or contact-sensitive axes."
        ),
        (
            "In a topology-preserving Synthetic Truth benchmark, the full model achieved AUROC 0.9919 and AUPRC 0.8333, whereas topology-anchor-only scoring retained high AUROC but dropped to AUPRC 0.5833. "
            "We analyzed the public 10x Genomics Preview Data: Atera In Situ Gene Expression, FFPE Human Breast Cancer dataset, generated using the pre-commercial Atera Whole Transcriptome Assay (Atera WTA) and Atera Onboard Analysis development workflow. "
            "TopoLink-CCI generated 1,319,600 common-resource hypotheses from 170,057 cells in this breast cancer tissue. "
            "The expanded benchmark terminalized 18 Breast WTA methods as nine full whole-dataset results and nine bounded subset results, with zero failure cards and zero deferred methods. "
            "Late-stage rescue runs converted FastCCC, SCILD, Copulacci and NicheNet into bounded appendix evidence, with NicheNet treated as downstream receiver-response support rather than a direct spatial ranker."
        ),
        (
            "TopoLink-CCI prioritized interpretable axes spanning vascular activation, stromal matrix biology, immune adhesion, Notch signaling and tumor-intrinsic adhesion. "
            "The top Breast WTA axis, VWF-SELP from endothelial cells to endothelial cells, is best interpreted as an endothelial activation and vascular adhesion niche consistent with Weibel-Palade body biology, not as direct proof of protein release. "
            "Cross-dataset application to the corresponding public 10x Genomics Preview Data: Atera In Situ Gene Expression, FFPE Human Cervical Cancer dataset produced 2,404,971 hypotheses from 717,576 cells and a distinct top tumor-adhesion axis, DSC2-DSG3 in differentiating tumor cells. "
            "TopoLink-CCI therefore prioritizes topology-supported spatial CCI hypotheses while leaving biochemical mechanism and causality to orthogonal experimental validation."
        ),
    ]
    figure_legends = [
        (
            "Figure 1 | TopoLink-CCI method and validation logic. "
            "a, Spatial CCI inference is vulnerable to co-expression, cell-abundance, proximity and resource-composition false positives. "
            "b, TopoLink-CCI scores each candidate using topology anchors, structure bridging, expression support and local contact. "
            "c, Workflow from Atera WTA preview datasets and pyXenium topology maps to ranked CCI axes and validation controls. "
            "d, Synthetic Truth evaluation shows AUROC 0.9919 and AUPRC 0.8333 for the full model, whereas topology-anchor-only scoring loses precision-recall performance. "
            "e, Orthogonal evidence matrix for seven interpretable breast cancer axes."
        ),
        (
            "Figure 2 | Whole-dataset benchmarking and biological interpretation. "
            "a, The expanded Breast WTA benchmark terminalizes 18 methods as nine full and nine bounded results. "
            "b, Interpretable Breast WTA axes ranked by TopoLink-CCI score are led by VWF-SELP. "
            "c, Canonical recovery is compared by within-method rank rather than raw score. "
            "d, Breast and cervical WTA datasets yield tissue-context-specific top axes. "
            "e, Bounded appendix methods provide scalability-aware terminal evidence."
        ),
    ]
    online_methods = [
        (
            "TopoLink-CCI score",
            "TopoLink-CCI evaluates each candidate molecular interaction axis using six retained components: sender topology anchor, receiver topology anchor, structure bridge, sender expression, receiver expression and local contact. The discovery score is computed as a prior-weighted geometric mean of these components.",
        ),
        (
            "Benchmark tiers",
            "Full results use whole-dataset common-resource outputs. Bounded results use documented cell-count, pair-count or method-specific scalability gates and are reported as appendix evidence. NicheNet is analyzed as downstream receiver-response support and is not presented as a direct spatial CCI ranker.",
        ),
        (
            "Validation controls",
            "Synthetic Truth metrics use defined positive and negative axes. Real WTA interpretation uses rank, canonical recovery, cell-label permutation, spatial nulls, matched-gene controls, downstream support, cross-method consistency, component ablation and bootstrap stability.",
        ),
    ]
    return {"abstract": abstract, "body": body, "figure_legends": figure_legends, "online_methods": online_methods}


def write_submission_docx() -> Path | None:
    if Document is None:
        return None
    sections = submission_text_sections()
    abstract = str(sections["abstract"])
    body = list(sections["body"])
    legends = list(sections["figure_legends"])
    main_package_words = word_count(abstract + " " + " ".join(body) + " " + " ".join(legends))
    if word_count(abstract) > 70:
        raise RuntimeError(f"Abstract exceeds 70 words: {word_count(abstract)}")
    if main_package_words > 1200:
        raise RuntimeError(f"Brief Communication main package exceeds 1200 words: {main_package_words}")

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    add_line_numbering(section)
    add_page_number(section)
    set_submission_styles(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.line_spacing = 2.0
    title.add_run("Topology-guided prioritization of spatial cell-cell interaction axes").bold = True
    add_double_spaced_paragraph(doc, "Author list: [To be completed]", bold=False)
    add_double_spaced_paragraph(doc, "Affiliations: [To be completed]", bold=False)
    doc.add_heading("Abstract", level=1)
    add_double_spaced_paragraph(doc, abstract)
    for paragraph in body:
        add_double_spaced_paragraph(doc, paragraph)

    doc.add_page_break()
    doc.add_heading("Figures and figure legends", level=1)
    for idx, (figure, legend) in enumerate(
        [
            ("figure_1_topolink_cci_method_validation.png", legends[0]),
            ("figure_2_topolink_cci_discovery_benchmark.png", legends[1]),
        ],
        start=1,
    ):
        path = FIG_DIR / figure
        if path.exists():
            doc.add_picture(str(path), width=Inches(6.2))
        add_double_spaced_paragraph(doc, legend)
        if idx == 1:
            doc.add_page_break()

    doc.add_page_break()
    doc.add_heading("Online Methods", level=1)
    for heading, text in sections["online_methods"]:
        doc.add_heading(str(heading), level=2)
        add_double_spaced_paragraph(doc, str(text))
    for heading, text in [
        ("Data availability", "The manuscript uses public 10x Genomics Preview Data: Atera In Situ Gene Expression, FFPE Human Breast Cancer and FFPE Human Cervical Cancer datasets generated using Atera Whole Transcriptome Assay (Atera WTA): https://www.10xgenomics.com/cn/datasets/atera-wta-ffpe-human-breast-cancer and https://www.10xgenomics.com/cn/datasets/atera-wta-ffpe-human-cervical-cancer. Benchmark outputs are archived in the repository artifact manifests."),
        ("Code availability", "TopoLink-CCI code and reproducible benchmark scripts are available in the pyXenium GitHub repository."),
        ("Acknowledgements", "[To be completed]"),
        ("Author contributions", "[To be completed]"),
        ("Competing interests", "The authors declare no competing interests."),
    ]:
        doc.add_heading(heading, level=1)
        add_double_spaced_paragraph(doc, text)

    out = MANUSCRIPT_DIR / "topolink_cci_short_communication_submission_initial.docx"
    doc.save(out)
    counts = pd.DataFrame(
        [
            {"item": "abstract", "word_count": word_count(abstract), "limit": 70, "status": "pass"},
            {"item": "abstract_body_legends", "word_count": main_package_words, "limit": 1200, "status": "pass"},
        ]
    )
    counts.to_csv(QA_DIR / "text_limit_qa.tsv", sep="\t", index=False)
    return out


def find_soffice() -> str | None:
    candidates = [
        shutil.which("soffice"),
        str(Path.home() / "scoop" / "apps" / "libreoffice" / "current" / "LibreOffice" / "program" / "soffice.exe"),
    ]
    for candidate in candidates:
        if candidate and Path(candidate).exists():
            return candidate
    return None


def convert_docx_to_pdf(docx_path: Path | None) -> Path | None:
    if docx_path is None:
        return None
    soffice = find_soffice()
    if soffice is None:
        return None
    pdf_path = docx_path.with_suffix(".pdf")
    if pdf_path.exists():
        pdf_path.unlink()
    subprocess.run(
        [
            soffice,
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(docx_path.parent),
            str(docx_path),
        ],
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
    )
    return pdf_path if pdf_path.exists() else None


def run_pdffonts(pdf_paths: list[Path]) -> Path:
    rows = []
    pdffonts = shutil.which("pdffonts")
    for pdf_path in pdf_paths:
        if not pdf_path or not pdf_path.exists():
            rows.append({"artifact": str(pdf_path), "status": "missing", "font_count": 0, "embedded_count": 0, "truetype_count": 0, "non_embedded_count": 0})
            continue
        if pdffonts is None:
            rows.append({"artifact": str(pdf_path), "status": "pdffonts_missing", "font_count": 0, "embedded_count": 0, "truetype_count": 0, "non_embedded_count": 0})
            continue
        result = subprocess.run([pdffonts, str(pdf_path)], check=False, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
        font_lines = [line for line in result.stdout.splitlines()[2:] if line.strip()]
        embedded_count = sum(1 for line in font_lines if " yes " in f" {line} ")
        truetype_count = sum(1 for line in font_lines if "TrueType" in line or "CID TrueType" in line)
        rows.append(
            {
                "artifact": str(pdf_path),
                "status": "pass" if result.returncode == 0 and font_lines else "check",
                "font_count": len(font_lines),
                "embedded_count": embedded_count,
                "truetype_count": truetype_count,
                "non_embedded_count": max(0, len(font_lines) - embedded_count),
                "pdffonts_stdout": result.stdout.replace("\n", "\\n"),
            }
        )
    path = QA_DIR / "font_qa_report.tsv"
    pd.DataFrame(rows).to_csv(path, sep="\t", index=False)
    return path


def write_figure_format_qa(outputs: dict[str, dict[str, str]]) -> Path:
    rows = []
    try:
        from PIL import Image
    except Exception:  # pragma: no cover
        Image = None
    for figure_id, fig_outputs in outputs.items():
        for ext, raw_path in fig_outputs.items():
            path = Path(raw_path)
            row = {
                "figure_id": figure_id,
                "format": ext,
                "path": str(path),
                "exists": path.exists(),
                "size_bytes": path.stat().st_size if path.exists() else 0,
                "rgb_or_vector": "vector" if ext in {"pdf", "svg"} else "not_checked",
                "dpi_x": "",
                "dpi_y": "",
                "width_px": "",
                "height_px": "",
                "nature_target": "180 mm width; max 170 mm height; text >=5 pt; RGB/300 dpi+ for raster",
            }
            if Image is not None and path.exists() and ext in {"png", "tiff"}:
                with Image.open(path) as image:
                    dpi = image.info.get("dpi", ("", ""))
                    row.update(
                        {
                            "rgb_or_vector": image.mode,
                            "dpi_x": round(float(dpi[0]), 2) if dpi and dpi[0] else "",
                            "dpi_y": round(float(dpi[1]), 2) if dpi and dpi[1] else "",
                            "width_px": image.width,
                            "height_px": image.height,
                        }
                    )
            rows.append(row)
    path = QA_DIR / "figure_format_qa.tsv"
    pd.DataFrame(rows).to_csv(path, sep="\t", index=False)
    return path


def render_pdf_previews(pdf_path: Path | None) -> list[Path]:
    if pdf_path is None or not pdf_path.exists() or shutil.which("pdftoppm") is None:
        return []
    out_prefix = QA_DIR / pdf_path.stem
    subprocess.run(["pdftoppm", "-png", "-r", "150", str(pdf_path), str(out_prefix)], check=True)
    return sorted(QA_DIR.glob(f"{pdf_path.stem}-*.png"))


def write_submission_readiness_checklist(submission_docx: Path | None, submission_pdf: Path | None, font_qa: Path, figure_qa: Path) -> Path:
    text = f"""# Nature Methods submission readiness checklist

- Target format: Nature Methods Brief Communication.
- Abstract limit: <=70 words; see `qa/text_limit_qa.tsv`.
- Main text package limit: <=1,200 words including abstract, references and figure legends; see `qa/text_limit_qa.tsv`.
- Main text structure: continuous main text without Results/Discussion-style subheadings; Online Methods retains subheadings.
- Display items: 2 main figures plus one fallback figure for editorial contingency.
- Manuscript DOCX: `{submission_docx}`.
- Manuscript PDF: `{submission_pdf}`.
- Font QA: `{font_qa}`.
- Figure format QA: `{figure_qa}`.
- Figure fonts: Matplotlib uses `pdf.fonttype=42` and `svg.fonttype=none`; verify with `pdffonts`.
- Figure raster exports: PNG/TIFF generated at 600 dpi; Nature minimum is 300 dpi.
- Figure source data: every panel has a TSV entry in `source_data/` and figure manifests in `metadata/`.
- Official checks consulted: Nature Methods content types, Nature initial submission guide and Nature figure specifications.
"""
    path = MANUSCRIPT_DIR / "submission_readiness_checklist.md"
    path.write_text(text, encoding="utf-8", newline="\n")
    return path


def write_docx(manuscript_paths: dict[str, Path]) -> Path | None:
    if Document is None:
        return None
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)
    for style_name, size in [("Normal", 9), ("Title", 16), ("Heading 1", 12), ("Heading 2", 10)]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
    main_text = manuscript_paths["main"].read_text(encoding="utf-8")
    for line in main_text.splitlines():
        if not line.strip():
            continue
        if line.startswith("# "):
            p = doc.add_paragraph(style="Title")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(line[2:]).bold = True
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.05
            p.add_run(line.replace("**", ""))
    doc.add_page_break()
    doc.add_heading("Main Figures", level=1)
    for figure in ["figure_1_topolink_cci_method_validation.png", "figure_2_topolink_cci_discovery_benchmark.png"]:
        path = FIG_DIR / figure
        if path.exists():
            doc.add_picture(str(path), width=Inches(6.9))
    out = MANUSCRIPT_DIR / "topolink_cci_short_communication_review_copy.docx"
    doc.save(out)
    return out


def write_package_manifest(
    outputs: dict[str, dict[str, str]],
    manuscripts: dict[str, Path],
    docx_path: Path | None,
    supplement: Path,
    submission_docx: Path | None,
    submission_pdf: Path | None,
    qa_artifacts: dict[str, Path],
) -> Path:
    review_pdf = docx_path.with_suffix(".pdf") if docx_path else None
    manifest = {
        "created_at": datetime.now(timezone.utc).isoformat(),
        "target": "Nature Methods Brief Communication",
        "fallback_target": "Nature Biotechnology Brief Communication",
        "backend": "Python/Matplotlib via nature-figure contract",
        "figures": outputs,
        "manuscripts": {key: str(path) for key, path in manuscripts.items()},
        "docx": str(docx_path) if docx_path else None,
        "review_pdf": str(review_pdf) if review_pdf and review_pdf.exists() else None,
        "submission_docx": str(submission_docx) if submission_docx else None,
        "submission_pdf": str(submission_pdf) if submission_pdf else None,
        "supplementary_manifest": str(supplement),
        "qa": {key: str(path) for key, path in qa_artifacts.items()},
    }
    path = PACKAGE_ROOT / "package_manifest.json"
    path.write_text(json.dumps(manifest, indent=2) + "\n", encoding="utf-8")
    return path


def validate_outputs(outputs: dict[str, dict[str, str]], required_paths: list[Path]) -> None:
    required = [
        ("Synthetic Truth AUROC", SYNTHETIC_TRUTH, "TopoLink-CCI"),
        ("method completion matrix", METHOD_MATRIX, "FastCCC"),
        ("VWF-SELP evidence", EVIDENCE, "VWF-SELP"),
    ]
    for label, path, token in required:
        text = path.read_text(encoding="utf-8", errors="replace")
        if token not in text:
            raise RuntimeError(f"Validation failed for {label}: missing {token}")
    for fig_outputs in outputs.values():
        for path in fig_outputs.values():
            if not Path(path).exists() or Path(path).stat().st_size == 0:
                raise RuntimeError(f"Missing figure output: {path}")
    for path in required_paths:
        if path is None or not path.exists() or path.stat().st_size == 0:
            raise RuntimeError(f"Missing required submission output: {path}")


def main() -> None:
    configure_matplotlib()
    ensure_dirs()
    data = load_data()
    outputs = {
        "figure_1": make_figure_1(data),
        "figure_2": make_figure_2(data),
        "one_figure_fallback": make_fallback(data),
    }
    manuscripts = write_manuscripts()
    supplement = write_supplement_manifest(data)
    docx_path = write_docx(manuscripts)
    review_pdf = convert_docx_to_pdf(docx_path)
    submission_docx = write_submission_docx()
    submission_pdf = convert_docx_to_pdf(submission_docx)
    pdfs_for_fonts = [Path(paths["pdf"]) for paths in outputs.values() if "pdf" in paths]
    for pdf_path in [review_pdf, submission_pdf]:
        if pdf_path is not None:
            pdfs_for_fonts.append(pdf_path)
    font_qa = run_pdffonts(pdfs_for_fonts)
    figure_qa = write_figure_format_qa(outputs)
    rendered_pages = render_pdf_previews(submission_pdf)
    checklist = write_submission_readiness_checklist(submission_docx, submission_pdf, font_qa, figure_qa)
    qa_artifacts = {
        "font_qa_report": font_qa,
        "figure_format_qa": figure_qa,
        "text_limit_qa": QA_DIR / "text_limit_qa.tsv",
        "submission_readiness_checklist": checklist,
    }
    if rendered_pages:
        qa_artifacts["submission_pdf_page_preview_1"] = rendered_pages[0]
    package_manifest = write_package_manifest(outputs, manuscripts, docx_path, supplement, submission_docx, submission_pdf, qa_artifacts)
    validate_outputs(outputs, [path for path in [submission_docx, submission_pdf, font_qa, figure_qa, checklist] if path is not None])
    print(json.dumps({"package": str(PACKAGE_ROOT), "manifest": str(package_manifest)}, indent=2))


if __name__ == "__main__":
    main()
