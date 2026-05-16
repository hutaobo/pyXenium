"""Build a one-main-figure short communication package for TopoLink-CCI.

The package is intentionally compact: one multipanel main figure, a concise
manuscript draft, figure source data, and a DOCX review copy.
"""

from __future__ import annotations

import json
import math
from datetime import datetime, timezone
from pathlib import Path

import matplotlib

matplotlib.use("Agg")

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from matplotlib import patches


try:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Inches, Pt
except Exception:  # pragma: no cover - DOCX is optional for figure generation
    Document = None


REPO = Path(__file__).resolve().parents[4]
CCI_ROOT = REPO / "benchmarking" / "cci_2026_atera"
SHORT_ROOT = CCI_ROOT / "topolink_cci_short_communication"
OUT = SHORT_ROOT / "one_main_figure"
FIG_DIR = OUT / "figures"
SOURCE_DIR = OUT / "source_data"
MANUSCRIPT_DIR = OUT / "manuscript"

METHOD_MATRIX = CCI_ROOT / "results" / "method_completion_matrix.tsv"
SYNTHETIC = (
    CCI_ROOT
    / "results"
    / "publication_benchmark_24h_20260511"
    / "pdc_collected"
    / "publication_benchmark_24h_20260511"
    / "synthetic_truth"
    / "tables"
    / "synthetic_truth_metrics.tsv"
)
EVIDENCE = CCI_ROOT / "pdc_validation_v2_collected" / "topolink_cci_validation_v2" / "tables" / "topolink_cci_validation_v2_evidence.tsv"
CROSS_DATASET = CCI_ROOT / "results" / "preview_completed_20260511" / "source_data" / "fig4_breast_cervical_topolink_comparison.tsv"


COLORS = {
    "teal": "#0F766E",
    "blue": "#2563EB",
    "orange": "#D97706",
    "red": "#B91C1C",
    "green": "#15803D",
    "purple": "#6D28D9",
    "gray": "#6B7280",
    "light": "#F3F4F6",
    "text": "#111827",
    "grid": "#E5E7EB",
    "bounded": "#FBBF24",
    "failure": "#FCA5A5",
    "full": "#86EFAC",
}

THEME_COLORS = {
    "vascular": "#0F766E",
    "stromal": "#8A5A2B",
    "immune": "#3F7F3F",
    "notch": "#5B5F97",
    "tumor": "#6B7280",
}


def configure_matplotlib() -> None:
    plt.rcParams.update(
        {
            "pdf.fonttype": 42,
            "ps.fonttype": 42,
            "svg.fonttype": "none",
            "font.family": "Arial",
            "font.size": 6,
            "axes.titlesize": 7,
            "axes.labelsize": 6,
            "xtick.labelsize": 5,
            "ytick.labelsize": 5,
            "legend.fontsize": 5,
            "axes.linewidth": 0.6,
            "lines.linewidth": 1.0,
            "savefig.dpi": 300,
        }
    )


def ensure_dirs() -> None:
    for path in [FIG_DIR, SOURCE_DIR, MANUSCRIPT_DIR]:
        path.mkdir(parents=True, exist_ok=True)


def load_inputs() -> dict[str, pd.DataFrame]:
    method = pd.read_csv(METHOD_MATRIX, sep="\t")
    synthetic = pd.read_csv(SYNTHETIC, sep="\t")
    evidence = pd.read_csv(EVIDENCE, sep="\t")
    evidence["axis"] = evidence["ligand"] + "-" + evidence["receptor"]
    cross = pd.read_csv(CROSS_DATASET, sep="\t")
    return {"method": method, "synthetic": synthetic, "evidence": evidence, "cross": cross}


def save_source_data(name: str, df: pd.DataFrame) -> Path:
    path = SOURCE_DIR / f"{name}.tsv"
    df.to_csv(path, sep="\t", index=False)
    return path


def add_panel_label(ax, label: str) -> None:
    ax.text(
        -0.08,
        1.08,
        label,
        transform=ax.transAxes,
        ha="left",
        va="top",
        fontsize=8,
        fontweight="bold",
        color=COLORS["text"],
    )


def rounded_box(ax, x: float, y: float, w: float, h: float, text: str, *, fc: str, ec: str, fs: float = 5.8) -> None:
    rect = patches.FancyBboxPatch(
        (x, y),
        w,
        h,
        boxstyle="round,pad=0.02,rounding_size=0.025",
        fc=fc,
        ec=ec,
        lw=0.8,
    )
    ax.add_patch(rect)
    ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=fs, color=COLORS["text"])


def panel_a(ax) -> pd.DataFrame:
    ax.set_axis_off()
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    add_panel_label(ax, "a")
    ax.set_title("Topology-guided CCI discovery", loc="left", fontweight="bold")
    nodes = [
        ("sender cell", 0.18, 0.58, COLORS["teal"]),
        ("receiver cell", 0.82, 0.58, COLORS["purple"]),
        ("tissue topology", 0.50, 0.82, COLORS["blue"]),
        ("local contact", 0.50, 0.28, COLORS["orange"]),
    ]
    for label, x, y, color in nodes:
        if "cell" in label:
            circ = patches.Circle((x, y), 0.075, fc=color, ec="white", lw=1.2)
            ax.add_patch(circ)
            ax.text(x, y, label.replace(" ", "\n"), ha="center", va="center", fontsize=5.5, color="white", fontweight="bold")
        else:
            rounded_box(ax, x - 0.14, y - 0.045, 0.28, 0.09, label, fc="#FFFFFF", ec=color)
    ax.annotate("", xy=(0.72, 0.58), xytext=(0.28, 0.58), arrowprops=dict(arrowstyle="->", lw=1.4, color=COLORS["text"]))
    ax.annotate("", xy=(0.50, 0.73), xytext=(0.25, 0.64), arrowprops=dict(arrowstyle="-", lw=0.8, color=COLORS["blue"]))
    ax.annotate("", xy=(0.50, 0.73), xytext=(0.75, 0.64), arrowprops=dict(arrowstyle="-", lw=0.8, color=COLORS["blue"]))
    ax.annotate("", xy=(0.50, 0.37), xytext=(0.25, 0.53), arrowprops=dict(arrowstyle="-", lw=0.8, color=COLORS["orange"]))
    ax.annotate("", xy=(0.50, 0.37), xytext=(0.75, 0.53), arrowprops=dict(arrowstyle="-", lw=0.8, color=COLORS["orange"]))
    ax.text(
        0.50,
        0.12,
        "CCI score = prior x geometric mean\n(sender anchor, receiver anchor, structure bridge,\nsender expression, receiver expression, local contact)",
        ha="center",
        va="center",
        fontsize=5.4,
        bbox=dict(fc="#F9FAFB", ec="#D1D5DB", boxstyle="round,pad=0.3"),
    )
    return pd.DataFrame(nodes, columns=["node", "x", "y", "color"])


def panel_b(ax, synthetic: pd.DataFrame) -> pd.DataFrame:
    add_panel_label(ax, "b")
    plot = synthetic.copy()
    plot["label"] = plot["method"].replace(
        {
            "TopoLink-CCI": "TopoLink-CCI",
            "expression_only": "expression",
            "contact_only": "contact",
            "topology_anchor_only": "topology\nanchor",
        }
    )
    x = np.arange(len(plot))
    width = 0.34
    ax.bar(x - width / 2, plot["auroc"], width=width, color=COLORS["blue"], label="AUROC")
    ax.bar(x + width / 2, plot["auprc"], width=width, color=COLORS["teal"], label="AUPRC")
    ax.set_ylim(0.45, 1.03)
    ax.set_xticks(x)
    ax.set_xticklabels(plot["label"], rotation=0)
    ax.set_title("Synthetic Truth validates full topology combination", loc="left", fontweight="bold")
    ax.set_ylabel("metric")
    ax.grid(axis="y", color=COLORS["grid"], lw=0.6)
    ax.legend(frameon=False, loc="lower left")
    top = plot.loc[plot["method"].eq("TopoLink-CCI")].iloc[0]
    ax.text(
        0.02,
        0.95,
        f"TopoLink-CCI\nAUROC {top.auroc:.4f}\nAUPRC {top.auprc:.4f}",
        transform=ax.transAxes,
        ha="left",
        va="top",
        fontsize=5.6,
        bbox=dict(fc="white", ec="#D1D5DB", boxstyle="round,pad=0.3"),
    )
    return plot[["method", "auroc", "auprc", "precision_at_5", "recall_at_5", "f1_at_5"]]


def panel_c(ax, evidence: pd.DataFrame) -> pd.DataFrame:
    add_panel_label(ax, "c")
    cols = [
        ("expr", "expression_specificity_support"),
        ("label", "cell_label_permutation_support"),
        ("spatial", "spatial_null_support"),
        ("matched", "matched_gene_control_support"),
        ("target", "downstream_target_support"),
        ("signal", "functional_received_signal_support"),
        ("cons.", "cross_method_support"),
        ("boot.", "bootstrap_stability_support"),
    ]
    df = evidence.sort_values("pyxenium_rank").head(7).copy()
    matrix = df[[c for _, c in cols]].astype(bool).to_numpy()
    ax.set_xlim(-0.5, len(cols) - 0.5)
    ax.set_ylim(len(df) - 0.5, -0.5)
    ax.set_xticks(range(len(cols)))
    ax.set_xticklabels([c for c, _ in cols], rotation=45, ha="right")
    ax.set_yticks(range(len(df)))
    ax.set_yticklabels(df["axis"].tolist())
    ax.set_title("Orthogonal controls support seven axes", loc="left", fontweight="bold")
    for i in range(matrix.shape[0]):
        for j in range(matrix.shape[1]):
            ax.scatter(
                j,
                i,
                s=58,
                color=COLORS["green"] if matrix[i, j] else COLORS["orange"],
                marker="o" if matrix[i, j] else "X",
                edgecolor="white",
                linewidth=0.5,
            )
    ax.grid(color=COLORS["grid"], lw=0.5)
    ax.tick_params(length=0)
    ax.text(0.98, 0.02, "7/7 strong\n0 artifact risk", transform=ax.transAxes, ha="right", va="bottom", fontsize=6, fontweight="bold")
    source = df[["axis", "pyxenium_rank", "CCI_score", "support_count", "evidence_class"] + [c for _, c in cols]]
    return source


def panel_d(ax, method: pd.DataFrame) -> pd.DataFrame:
    add_panel_label(ax, "d")
    breast = method.loc[method["dataset"].eq("atera_breast_wta")].copy()
    order = ["full_result", "bounded_subset_result", "reproducible_failure_card"]
    counts = breast["status"].value_counts().reindex(order).fillna(0).astype(int)
    labels = ["full", "bounded", "failure\ncard"]
    colors = [COLORS["full"], COLORS["bounded"], COLORS["failure"]]
    ax.bar(labels, counts.values, color=colors, edgecolor="#374151", linewidth=0.6)
    ax.set_ylim(0, max(counts.values) + 2)
    ax.set_ylabel("methods")
    ax.set_title("Expanded benchmark is terminalized", loc="left", fontweight="bold")
    ax.grid(axis="y", color=COLORS["grid"], lw=0.6)
    for i, v in enumerate(counts.values):
        ax.text(i, v + 0.25, str(v), ha="center", va="bottom", fontsize=7, fontweight="bold")
    ax.text(0.98, 0.94, "18 breast methods\n0 deferred", transform=ax.transAxes, ha="right", va="top", fontsize=6)
    return counts.rename_axis("status").reset_index(name="n_methods")


def panel_e(ax, cross: pd.DataFrame) -> pd.DataFrame:
    add_panel_label(ax, "e")
    df = cross.copy()
    top = df.groupby("dataset", group_keys=False).head(5).copy()
    # Preserve displayed order by dataset and score.
    top["display"] = top["dataset"].str.title() + ": " + top["axis"]
    top = top.sort_values(["dataset", "CCI_score"], ascending=[True, True])
    colors = [COLORS["teal"] if x == "breast" else COLORS["purple"] for x in top["dataset"]]
    ax.barh(top["display"], top["CCI_score"], color=colors)
    ax.set_xlim(0.60, max(0.84, float(top["CCI_score"].max()) + 0.02))
    ax.set_xlabel("CCI score")
    ax.set_title("Whole-dataset discovery generalizes across tissues", loc="left", fontweight="bold")
    ax.grid(axis="x", color=COLORS["grid"], lw=0.6)
    ax.text(
        0.02,
        0.03,
        "Breast top: VWF-SELP endothelial activation\nCervical top: DSC2-DSG3 differentiating tumor adhesion",
        transform=ax.transAxes,
        ha="left",
        va="bottom",
        fontsize=5.4,
        bbox=dict(fc="white", ec="#D1D5DB", boxstyle="round,pad=0.25"),
    )
    return top[["dataset", "axis", "sender_receiver", "CCI_score", "local_contact", "cross_edge_count"]]


def build_figure(data: dict[str, pd.DataFrame]) -> dict[str, Path]:
    configure_matplotlib()
    fig = plt.figure(figsize=(7.2, 8.6), constrained_layout=True)
    gs = fig.add_gridspec(3, 2, height_ratios=[1.05, 1.15, 1.05], width_ratios=[1.0, 1.25])

    source_paths = {}
    source_paths["panel_a_schematic"] = save_source_data("figure1_panel_a_schematic", panel_a(fig.add_subplot(gs[0, 0])))
    source_paths["panel_b_synthetic_truth"] = save_source_data("figure1_panel_b_synthetic_truth", panel_b(fig.add_subplot(gs[0, 1]), data["synthetic"]))
    source_paths["panel_c_evidence"] = save_source_data("figure1_panel_c_evidence_matrix", panel_c(fig.add_subplot(gs[1, :]), data["evidence"]))
    source_paths["panel_d_status"] = save_source_data("figure1_panel_d_method_status", panel_d(fig.add_subplot(gs[2, 0]), data["method"]))
    source_paths["panel_e_cross_dataset"] = save_source_data("figure1_panel_e_cross_dataset_axes", panel_e(fig.add_subplot(gs[2, 1]), data["cross"]))

    fig.suptitle("TopoLink-CCI prioritizes topology-supported spatial cell-cell interaction hypotheses", fontsize=9, fontweight="bold")

    outputs = {}
    for ext in ["pdf", "svg", "png"]:
        path = FIG_DIR / f"figure_1_topolink_cci_one_main.{ext}"
        fig.savefig(path, bbox_inches="tight", dpi=300)
        outputs[ext] = path
    plt.close(fig)

    manifest = {
        "created_at": datetime.now(timezone.utc).isoformat(),
        "figure_outputs": {k: str(v) for k, v in outputs.items()},
        "source_data": {k: str(v) for k, v in source_paths.items()},
        "style": {"pdf.fonttype": 42, "svg.fonttype": "none", "font": "Arial"},
    }
    (OUT / "figure_manifest.json").write_text(json.dumps(manifest, indent=2) + "\n", encoding="utf-8")
    return outputs


def manuscript_text() -> str:
    return """# Topology-guided prioritization of spatial cell-cell interaction axes

**Target format:** Nature Methods Brief Communication or Nature Biotechnology Brief Communication
**Preferred target:** Nature Methods Brief Communication
**Display item strategy:** one main figure plus supplementary data and online methods
**Method name:** TopoLink-CCI

## Editorial positioning

Nature Methods is the stronger first target because the manuscript is a method and validation framework for spatial omics CCI inference. Nature Biotechnology is possible only if the emphasis is shifted toward broad enabling value for translational spatial tissue analysis; its Brief Communication page also notes that the format is not intended for presentation of research data or analysis alone. This draft therefore frames TopoLink-CCI as a concise method/tool report, with the benchmark used as validation rather than as the sole claim.

## Abstract

Spatial cell-cell interaction inference is confounded by co-expression, abundance and incomplete molecular resources. We introduce TopoLink-CCI, a topology-guided framework that combines tissue topology, expression specificity and local contact, then challenges candidates with false-positive controls. Across Xenium WTA breast and cervical cancers, TopoLink-CCI scales to whole datasets and prioritizes interpretable vascular, stromal, immune and tumor interaction axes.

## Main text draft

Spatial transcriptomics has made it possible to infer cell-cell interaction (CCI) axes in intact tissues, but high-resolution whole-transcriptome imaging also increases false-positive risk. A ligand and receptor can appear biologically plausible because both are strongly expressed, because the corresponding cell types are abundant, or because a curated molecular resource includes extracellular-matrix, adhesion, scavenger-receptor and shared activation-state relationships that are not classical secreted signaling events. Spatial proximity improves interpretation but does not by itself prove molecular exchange, receptor activation or causal signaling. The central problem is therefore not only to rank candidate ligand-receptor-resource pairs, but to prioritize CCI hypotheses that are compatible with tissue topology, local organization and independent controls.

We developed TopoLink-CCI as a topology-guided CCI prioritization framework in pyXenium. In CCI-resource mode, each ligand, receptor, sender and receiver combination is scored by six retained components: sender topology anchor, receiver topology anchor, sender-receiver structure bridge, sender expression support, receiver expression support and local contact support. The discovery score is a prior-weighted geometric mean of these components, so a high score requires concordance across topology, expression and spatial contact rather than any single high-expression feature. Because every component is exported, users can diagnose whether an axis is topology-driven, expression-driven, contact-sensitive or potentially dominated by a resource prior. This design explicitly treats the score as a discovery statistic, not as proof of protein-level communication.

We first evaluated the scoring logic in a topology-preserving Synthetic Truth benchmark. The full TopoLink-CCI model achieved AUROC 0.9919 and AUPRC 0.8333, whereas the topology-anchor-only model retained high AUROC (0.9839) but dropped to AUPRC 0.5833. This difference shows that topology anchors alone are insufficient for stable precision-recall ranking and supports the combined use of expression specificity and local contact. We then applied TopoLink-CCI to Atera Xenium WTA breast cancer, generating 1,319,600 common-resource CCI hypotheses from 170,057 cells. The expanded breast benchmark now includes 18 terminalized methods: nine full results, six bounded subset results and three reproducible failure cards, with no deferred methods. Comparison methods included CellPhoneDB, LARIS, LIANA+, SpatialDM, stLearn, Squidpy, CellChat, COMMOT, Giotto, SpaTalk, NICHES, CellNEST, CellAgentChat and FastCCC, with bounded results separated from full whole-dataset results.

TopoLink-CCI prioritized biologically interpretable axes spanning vascular activation, stromal matrix biology, immune adhesion, Notch signaling and tumor-intrinsic adhesion. Seven representative axes were evaluated with orthogonal controls inspired by established CCI methods, including cell-label permutation, spatial nulls, matched-expression gene controls, downstream target support, received-signal association, cross-method consensus, component ablation and bootstrap stability. All seven were classified as strong computational hypotheses with no contamination flag, although the evidence matrix preserved axis-specific caveats rather than enforcing a single pass-fail criterion. The top breast axis, VWF-SELP from endothelial cells to endothelial cells, is best interpreted as an endothelial activation and vascular adhesion niche consistent with Weibel-Palade body biology, not as direct proof of VWF/P-selectin protein release. Cross-dataset application to Xenium WTA cervical cancer produced 2,404,971 hypotheses and a distinct top tumor-adhesion axis, DSC2-DSG3 in differentiating tumor cells, supporting tissue-context-specific prioritization.

These results position TopoLink-CCI as a spatial CCI hypothesis-prioritization framework rather than a universal proof engine for intercellular signaling. Its value is to reduce a large molecular search space to axes that are consistent with tissue topology, expression specificity, local contact and independent computational controls. The framework is designed to guide downstream mechanistic validation, including protein-level spatial assays, perturbation experiments or targeted spatial multi-omics, while keeping computational evidence separate from causal biological proof.

## Main figure legend

**Figure 1 | TopoLink-CCI prioritizes topology-supported spatial CCI hypotheses.** **a,** TopoLink-CCI integrates sender and receiver topology anchors, tissue-structure bridging and local cell-cell contact. **b,** Synthetic Truth evaluation shows that the full model maintains high AUROC and AUPRC, whereas topology-anchor-only scoring loses precision-recall performance. **c,** Orthogonal evidence matrix for seven interpretable breast cancer axes. Green circles indicate support and orange crosses indicate axis-specific weak evidence layers. **d,** Terminal expanded breast benchmark status: nine full results, six bounded subset results and three reproducible failure cards. **e,** Whole-dataset TopoLink-CCI discoveries differ between breast and cervical WTA datasets, supporting tissue-context-specific prioritization.

## Required supplementary package

- Supplementary Note 1: complete score definition and component interpretation.
- Supplementary Note 2: synthetic truth generation and false-positive controls.
- Supplementary Table 1: full method completion matrix.
- Supplementary Table 2: validation evidence table for seven axes.
- Supplementary Table 3: runtime, memory and scalability summary.
- Supplementary Data: standardized CCI outputs or indexed manifests, not bulky raw caches.

## Guardrails

- Do not claim that TopoLink-CCI proves protein binding, secretion, receptor activation or causal signaling.
- Do not compare raw scores across methods.
- Use F1/AUROC/AUPRC only where ground truth or predefined canonical axes exist.
- Keep bounded subset methods distinct from full whole-dataset methods.
"""


def write_manuscript_docx(markdown: str, figure_png: Path) -> Path | None:
    if Document is None:
        return None
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)

    styles = doc.styles
    for name, size, bold in [("Normal", 9, False), ("Title", 16, True), ("Heading 1", 12, True), ("Heading 2", 10, True)]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = bold

    for line in markdown.splitlines():
        if not line.strip():
            continue
        if line.startswith("# "):
            p = doc.add_paragraph(style="Title")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(line[2:]).bold = True
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.05
            p.add_run(line.replace("**", ""))

    doc.add_page_break()
    doc.add_heading("Main Figure", level=1)
    doc.add_picture(str(figure_png), width=Inches(6.9))
    out = MANUSCRIPT_DIR / "topolink_cci_one_main_figure_short_communication.docx"
    doc.save(out)
    return out


def main() -> None:
    ensure_dirs()
    data = load_inputs()
    outputs = build_figure(data)
    md = manuscript_text()
    md_path = MANUSCRIPT_DIR / "topolink_cci_one_main_figure_short_communication.md"
    md_path.write_text(md, encoding="utf-8", newline="\n")
    docx_path = write_manuscript_docx(md, outputs["png"])
    package = {
        "created_at": datetime.now(timezone.utc).isoformat(),
        "manuscript_md": str(md_path),
        "manuscript_docx": str(docx_path) if docx_path else None,
        "figure": {k: str(v) for k, v in outputs.items()},
        "source_data_dir": str(SOURCE_DIR),
        "journal_fit": {
            "preferred": "Nature Methods Brief Communication",
            "secondary": "Nature Biotechnology Brief Communication if positioned as a broadly enabling spatial tissue-analysis method",
        },
    }
    (OUT / "package_manifest.json").write_text(json.dumps(package, indent=2) + "\n", encoding="utf-8")
    print(json.dumps(package, indent=2))


if __name__ == "__main__":
    main()
